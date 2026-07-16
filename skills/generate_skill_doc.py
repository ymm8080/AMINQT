# -*- coding: utf-8 -*-
"""Generate a Word document listing all installed skills with invocation samples."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)

# ── Helper functions ──
def add_code_block(doc, code_text, font_size=9):
    """Add a code block with monospace formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    # Shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    p._p.get_or_add_pPr().append(shading)

def add_heading_styled(doc, text, level=1):
    """Add a heading with custom styling."""
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False, italic=False, size=11):
    """Add a normal paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(text)
        run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)

# ════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════
title = doc.add_heading('AMINQT Skills Reference', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('How to Invoke Each Installed Skill')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
run.italic = True

doc.add_paragraph()

# Overview table
add_heading_styled(doc, 'Installed Skills Overview', level=1)

overview_data = [
    ['Skill Name', 'Type', 'Source', 'Invocation'],
    ['quant-ohlcv-feature', 'Python Library', 'YuxinSUN89/quant-ohlcv-feature', 'import + signal(df, n, name)'],
    ['ohlcpattern', 'Python Library', 'tempusoneps/ohlcpattern', 'CandlestickPatterns(df)'],
    ['aurumq-rl', 'CLI + Python', 'yupoet/aurumq-rl', 'aurumq-rl-infer / from aurumq_rl import'],
    ['quantitative_strategy_mcp', 'Python Library', 'Skywalkerhm/quantitative_strategy_mcp', 'call_mcp_tool()'],
    ['investagent', 'AI Skill (SKILL.md)', 'tohnee/investagent', 'Chat trigger phrases'],
    ['a-share-selection-strategy', 'AI Skill (SKILL.md)', 'MisonL/a-share-selection-strategy-skill', 'Chat trigger phrases'],
    ['uzi-skill (sub-module)', 'AI Skill + Script', 'wbh604/uzi-skill (via investagent)', 'python run.py <ticker>'],
    ['buffett-skills (sub-module)', 'AI Skill', 'agi-now/buffett-skills (via investagent)', 'Chat: "Buffett framework"'],
    ['serenity-skill (sub-module)', 'AI Skill', 'via investagent', 'Chat: "industry chain"'],
    ['TradingAgents (sub-module)', 'Python Script', 'TauricResearch/TradingAgents (via investagent)', 'python main.py --ticker'],
    ['QuantDinger (sub-module)', 'Docker + API', 'brokermr810/QuantDinger (via investagent)', 'docker-compose + API'],
]

table = doc.add_table(rows=len(overview_data), cols=4)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(overview_data):
    for j, cell_text in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. quant-ohlcv-feature
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, '1. quant-ohlcv-feature', level=1)
add_para(doc, '335+ technical indicator implementations for OHLCV data. Each indicator is a standalone function with a uniform signal(df, n, factor_name) signature.')

add_heading_styled(doc, 'Categories', level=2)
add_bullet(doc, '121 files — RSI, MACD, KDJ, CCI, CMO, ROC, MTM, DBCD, Fisher', bold_prefix='momentum_feature/ — ')
add_bullet(doc, '67 files — ADX, DEMA, TEMA, HMA, KAMA, Ichimoku, Aroon, BBI', bold_prefix='trend_feature/ — ')
add_bullet(doc, '72 files — Bollinger, ATR, Keltner, Donchian, Fibonacci, APZ', bold_prefix='volatility_feature/ — ')
add_bullet(doc, '51 files — OBV, CMF, MFI, Force Index, EMV, PVT, WVAD', bold_prefix='volume_feature/ — ')
add_bullet(doc, '9 files — VWAP, Typical Price, Weighted Close, AvgPrice', bold_prefix='price_feature/ — ')
add_bullet(doc, '4 files — Amihud illiquidity, Bid-Ask Spread, MarketPL', bold_prefix='liquidity_feature/ — ')
add_bullet(doc, '11 files — FearGreed, DAMAO, MSBT, CoppAtrBull', bold_prefix='composite_feature/ — ')

add_heading_styled(doc, 'Sample: Single Indicator', level=2)
add_code_block(doc, '''import pandas as pd
from momentum_feature.Rsi import signal as rsi_signal

df = pd.read_csv("data/raw/600519.csv")
df = rsi_signal(df, n=14, factor_name="rsi_14")

print(df[["close", "rsi_14"]].tail())''')

add_heading_styled(doc, 'Sample: Multi-Indicator Pipeline', level=2)
add_code_block(doc, '''import pandas as pd
from momentum_feature.Rsi import signal as rsi
from momentum_feature.Macd import signal as macd
from trend_feature.Adx import signal as adx
from volatility_feature.Bolling import signal as boll
from volume_feature.Obv import signal as obv

df = pd.read_csv("data/raw/600519.csv")

# All indicators share the same signature
df = rsi(df, n=14, factor_name="rsi_14")
df = macd(df, n=12, factor_name="macd_12")
df = adx(df, n=14, factor_name="adx_14")
df = boll(df, n=20, factor_name="boll_20")
df = obv(df, n=14, factor_name="obv_14")

print(df[["close", "rsi_14", "macd_12", "adx_14", "boll_20", "obv_14"]].tail())''')

add_heading_styled(doc, 'Sample: Batch Process All Indicators in a Category', level=2)
add_code_block(doc, '''import pandas as pd
import importlib
import os
import glob

df = pd.read_csv("data/raw/600519.csv")

# Auto-import and run all momentum features
feature_dir = "AMINQT CODES/skills/quant-ohlcv-feature/momentum_feature"
for pyfile in glob.glob(os.path.join(feature_dir, "*.py")):
    modname = os.path.basename(pyfile).replace(".py", "")
    if modname.startswith("_"):
        continue
    try:
        mod = importlib.import_module(f"momentum_feature.{modname}")
        df = mod.signal(df, n=14, factor_name=modname.lower())
    except Exception as e:
        print(f"SKIP {modname}: {e}")

print(f"Total columns: {len(df.columns)}")''')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 2. ohlcpattern
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, '2. ohlcpattern', level=1)
add_para(doc, 'Detects 40+ candlestick patterns in financial data. Built with pandas for speed and ease of use.')

add_heading_styled(doc, 'Prerequisites', level=2)
add_para(doc, 'DataFrame must contain columns: Open, High, Low, Close (capitalized first letter).')

add_heading_styled(doc, 'Sample: Basic Pattern Detection', level=2)
add_code_block(doc, '''import pandas as pd
from ohlcpattern.candlestick import CandlestickPatterns

# 1. Load data (ensure columns: Open, High, Low, Close)
df = pd.read_csv('data/raw/600519.csv')
df.columns = [c.capitalize() for c in df.columns]  # open -> Open

# 2. Initialize pattern detector
csp = CandlestickPatterns(df)

# 3. Add detection category: 'reversal', 'continuation', or 'full'
csp._add('reversal')

# 4. Generate modeling result
modeling_data = csp.pattern_modeling()

# 5. Filter for detected patterns
detected = modeling_data[modeling_data.model != '']
print(detected[['Open', 'High', 'Low', 'Close', 'model']])''')

add_heading_styled(doc, 'Sample: Full Pattern Scan (All Categories)', level=2)
add_code_block(doc, '''import pandas as pd
from ohlcpattern.candlestick import CandlestickPatterns

df = pd.read_csv('data/raw/600519.csv')
df.columns = [c.capitalize() for c in df.columns]

csp = CandlestickPatterns(df)
csp._add('full')  # reversal + continuation
modeling_data = csp.pattern_modeling()

# Count patterns by type
pattern_counts = modeling_data[modeling_data.model != '']['model'].value_counts()
print("Detected patterns:")
print(pattern_counts)

# Export results
modeling_data.to_csv('output/candlestick_patterns.csv', index=False)''')

add_heading_styled(doc, 'CLI Usage', level=2)
add_code_block(doc, '''# Extract patterns from CSV file
ohlcpattern extract data/raw/600519.csv --output patterns.csv

# Show help
ohlcpattern extract --help

# Check version
ohlcpattern --version''')

add_heading_styled(doc, 'Supported Patterns', level=2)
patterns_data = [
    ['Type', 'Bullish', 'Bearish'],
    ['Single', 'Hammer, Inverted Hammer', 'Shooting Star, Hanging Man'],
    ['Double', 'Bullish Engulfing, Piercing, Bullish Harami, Tweezers Bottom', 'Bearish Engulfing, Dark Cloud, Bearish Harami, Tweezers Top'],
    ['Triple', 'Morning Star, Three White Soldiers', 'Evening Star, Three Black Crows'],
    ['Continuation', 'Rising Three Methods, Bullish Gap, Fair Value Gap', 'Falling Three Methods, Bearish Gap, Neck Pattern'],
]
ptable = doc.add_table(rows=len(patterns_data), cols=3)
ptable.style = 'Light Shading Accent 1'
for i, row_data in enumerate(patterns_data):
    for j, cell_text in enumerate(row_data):
        cell = ptable.rows[i].cells[j]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. aurumq-rl
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, '3. aurumq-rl', level=1)
add_para(doc, 'Reinforcement learning stock-selection framework for China A-share market. 296 price-volume factors + PPO/A2C/SAC training + ONNX inference pipeline.')

add_heading_styled(doc, 'CLI Commands', level=2)
add_code_block(doc, '''# Train a model (requires GPU + train dependencies)
aurumq-rl-train --config configs/default.yaml

# Run inference with a trained model (CPU is fine)
aurumq-rl-infer --model-path models/best.onnx --data data/raw/

# Export factor panel to parquet
aurumq-rl-export --output factor_panel.parquet

# CPU-only smoke test (no GPU needed, ~30 seconds)
cd "AMINQT CODES/skills/aurumq-rl"
python examples/quickstart.py

# Run tests
cd "AMINQT CODES/skills/aurumq-rl"
python -m pytest tests/ -v -m smoke''')

add_heading_styled(doc, 'Sample: Python API — Quickstart', level=2)
add_code_block(doc, '''from aurumq_rl import FactorPanelLoader
from aurumq_rl.data_loader import FACTOR_COL_PREFIXES

# Build a synthetic factor panel for testing
panel = FactorPanelLoader.build_synthetic(
    n_dates=60, n_stocks=50, n_factors=12,
    forward_period=5, seed=42
)
print(f"Factor array shape: {panel.factor_array.shape}")
print(f"Return array shape: {panel.return_array.shape}")
print(f"Dates: {panel.dates[0]} .. {panel.dates[-1]}")
print(f"Factor names: {panel.factor_names[:3]}...")''')

add_heading_styled(doc, 'Sample: Run RL Environment', level=2)
add_code_block(doc, '''import datetime
import numpy as np
from aurumq_rl import FactorPanelLoader
from aurumq_rl.env import StockPickingConfig, StockPickingEnv

# Build panel
panel = FactorPanelLoader.build_synthetic(
    n_dates=60, n_stocks=50, n_factors=12,
    forward_period=5, seed=42
)

# Configure environment
config = StockPickingConfig(
    start_date=datetime.date(2022, 1, 1),
    end_date=datetime.date(2022, 12, 31),
    n_factors=12,
    top_k=10,
    forward_period=5,
    cost_bps=20.0,
    turnover_penalty=0.0,
)

# Create environment
env = StockPickingEnv(
    config=config,
    factor_panel=panel.factor_array,
    return_panel=panel.return_array,
    pct_change_panel=panel.pct_change_array,
    is_st_panel=panel.is_st_array,
    is_suspended_panel=panel.is_suspended_array,
    days_since_ipo_panel=panel.days_since_ipo_array,
)

# Run steps
obs, _info = env.reset(seed=0)
for step in range(20):
    action = np.random.uniform(0, 1, size=50).astype(np.float32)
    obs, reward, terminated, truncated, info = env.step(action)
    print(f"Step {step}: reward={reward:+.5f}")
    if terminated or truncated:
        break''')

add_heading_styled(doc, 'Sample: Inference with ONNX', level=2)
add_code_block(doc, '''from aurumq_rl.scripts.infer import run_inference

results = run_inference(
    model_path="models/best.onnx",
    data_dir="data/raw/"
)
print(f"Selected stocks: {results['selected']}")
print(f"Scores: {results['scores']}")''')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. quantitative_strategy_mcp
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, '4. quantitative_strategy_mcp', level=1)
add_para(doc, 'MCP-based quantitative strategy toolkit: stock data retrieval, factor calculation, backtesting, and stock monitoring.')

add_heading_styled(doc, 'Setup', level=2)
add_code_block(doc, '''import sys, os
sys.path.insert(0, os.path.join(
    "AMINQT CODES", "skills", "quantitative_strategy_mcp", "src"
))
os.environ['DB_PATH'] = '/path/to/your/stock_database.db'

from mcp_tools import call_mcp_tool, list_all_tools

# List all available tools
print(list_all_tools())''')

add_heading_styled(doc, 'Sample: Get Stock Data', level=2)
add_code_block(doc, '''from mcp_tools import call_mcp_tool

result = call_mcp_tool(
    'get_stock_data',
    ts_code='000001.SZ',
    start_date='20240101',
    end_date='20241231'
)

if result['success']:
    print(f"Got {result['count']} records")
    print(f"First record: {result['data'][0]}")
else:
    print(f"Error: {result['message']}")''')

add_heading_styled(doc, 'Sample: Calculate Factor (RSI)', level=2)
add_code_block(doc, '''from mcp_tools import call_mcp_tool

# Step 1: Get data
data_result = call_mcp_tool(
    'get_stock_data',
    ts_code='000001.SZ',
    start_date='20240101',
    end_date='20241231'
)

# Step 2: Calculate RSI factor
factor_result = call_mcp_tool(
    'calculate_factor',
    data=data_result['data'],
    factor_type='rsi',
    params={'period': 14}
)

if factor_result['success']:
    print(f"Latest RSI: {factor_result['factor_values'][-1]:.2f}")''')

add_heading_styled(doc, 'Sample: Run Backtest', level=2)
add_code_block(doc, '''from mcp_tools import call_mcp_tool

result = call_mcp_tool(
    'run_backtest',
    ts_code='000001.SZ',
    start_date='20240101',
    end_date='20241231',
    strategy_type='momentum',
    params={'lookback': 20, 'threshold': 0.02}
)

if result['success']:
    m = result['metrics']
    print(f"Total Return:  {m['total_return']*100:.1f}%")
    print(f"Annual Return: {m['annual_return']*100:.1f}%")
    print(f"Sharpe Ratio:  {m['sharpe_ratio']:.2f}")
    print(f"Max Drawdown:  {m['max_drawdown']*100:.1f}%")
    print(f"Win Rate:      {m['win_rate']*100:.1f}%")''')

add_heading_styled(doc, 'Sample: Monitor Stocks', level=2)
add_code_block(doc, '''from mcp_tools import call_mcp_tool

result = call_mcp_tool(
    'monitor_stocks',
    stock_list=['000001.SZ', '000002.SZ', '600036.SH'],
    threshold=0.03  # 3% alert threshold
)

if result['success']:
    if result['alerts']:
        for alert in result['alerts']:
            print(f"ALERT: {alert['message']}")
    else:
        print("No alerts")''')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5. investagent (AI Skill with 5 sub-frameworks)
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, '5. investagent (Full-Stack AI Investment Agent)', level=1)
add_para(doc, 'Integrates 5 research frameworks. Invoke by saying trigger phrases in chat. The AI reads SKILL.md and routes to the correct framework automatically.')

add_heading_styled(doc, 'Invocation: Chat Trigger Phrases', level=2)
trigger_data = [
    ['Trigger Phrase', 'Framework Activated', 'Output'],
    ['"Research A-share AI semiconductor"', 'Serenity -> Buffett -> UZI -> TradingAgents -> QuantDinger', 'Full pipeline report'],
    ['"Deep analysis of Kweichow Moutai"', 'Buffett -> UZI -> TradingAgents -> QuantDinger', 'Stock deep-dive report'],
    ['"Analyze CPO supply chain bottlenecks"', 'Serenity (full workflow)', 'Industry chain map'],
    ['"Buffett framework analysis of Apple"', 'Buffett (full analysis)', 'Value investing report'],
    ['"Multi-agent analysis of NVDA"', 'TradingAgents', 'Multi-agent decision'],
    ['"Backtest a dual-MA strategy on BTC"', 'QuantDinger', 'Backtest PnL curve'],
]
ttable = doc.add_table(rows=len(trigger_data), cols=3)
ttable.style = 'Light Shading Accent 1'
for i, row_data in enumerate(trigger_data):
    for j, cell_text in enumerate(row_data):
        cell = ttable.rows[i].cells[j]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

add_heading_styled(doc, 'Sub-Module: UZI-Skill (22-dim Deep Analysis)', level=2)
add_para(doc, '22-dimensional data collection + 65-person panel scoring + 6 valuation models. Produces Bloomberg-style HTML report.')
add_code_block(doc, '''# Command-line invocation
cd "AMINQT CODES/skills/investagent/skills/uzi-skill"
python run.py 600519.SH

# Output is saved to:
#   .cache/600519.SH/  (JSON data files)
#   reports/           (HTML report)''')

add_heading_styled(doc, 'Sub-Module: Buffett Skills (Value Investing)', level=2)
add_para(doc, '8-question quick screening + 10-chapter deep analysis. Invoked via chat or SKILL.md.')
add_code_block(doc, '''# AI reads the SKILL.md and follows the protocol:
# 1. Read: skills/buffett-skills/skills/buffett/SKILL.md
# 2. Execute 8-question quick screening
# 3. If passed, run 10-chapter deep analysis
# 4. Output: structured value investing report

# Chat trigger:
#   "Use Buffett framework to analyze Apple"
#   "Buffett screening for 600519"''')

add_heading_styled(doc, 'Sub-Module: Serenity Skill (Industry Chain)', level=2)
add_code_block(doc, '''# AI reads the SKILL.md and follows the workflow:
# 1. Read: skills/serenity-skill/SKILL.md
# 2. Read: skills/serenity-skill/references/deep-research-workflow.md
# 3. Execute 9-step research workflow
# 4. Output: 8-layer value chain map + scarce layers + 20+ company pool

# Chat trigger:
#   "Analyze CPO supply chain bottlenecks"
#   "Research A-share AI semiconductor sector"''')

add_heading_styled(doc, 'Sub-Module: TradingAgents (Multi-Agent Decision)', level=2)
add_code_block(doc, '''# Requires LLM API key
cd "AMINQT CODES/skills/investagent/skills/TradingAgents"
pip install -r requirements.txt

# Run multi-agent analysis
python main.py --ticker NVDA --date 2024-01-15

# Chat trigger:
#   "Multi-agent analysis of NVDA"
#   "Bull vs bear debate on Kweichow Moutai"''')

add_heading_styled(doc, 'Sub-Module: QuantDinger (Backtest Engine)', level=2)
add_code_block(doc, '''# Docker-based backtest engine
cd "AMINQT CODES/skills/investagent/skills/QuantDinger"
docker-compose up -d

# Chat trigger:
#   "Backtest a dual-MA strategy on BTC"
#   "Write a momentum strategy and test it"''')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. a-share-selection-strategy
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, '6. a-share-selection-strategy', level=1)
add_para(doc, 'Reproducible, auditable A-share stock selection workflow. AI reads SKILL.md and routes to the correct task path.')

add_heading_styled(doc, 'Invocation: Chat Trigger Phrases', level=2)
trigger2 = [
    ['Trigger Phrase', 'Task Path', 'Action'],
    ['"Today A-share selection"', 'Full-A strict task', 'Scan all A-shares, score, filter, rank'],
    ['"Select low-price stocks"', 'Low-price short-term', 'Screen by price thresholds, quick scoring'],
    ['"Expand stock pool scan"', 'Full-market scan', 'Broad market coverage with provenance'],
    ['"Local scoring with prices.csv"', 'Local scoring', 'Validate OHLCV, score candidates locally'],
    ['"Prediction-derived scoring"', 'Prediction-derived', 'Use external predictions for scoring'],
]
t2table = doc.add_table(rows=len(trigger2), cols=3)
t2table.style = 'Light Shading Accent 1'
for i, row_data in enumerate(trigger2):
    for j, cell_text in enumerate(row_data):
        cell = t2table.rows[i].cells[j]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

add_heading_styled(doc, 'Sample: Run Selection Script', level=2)
add_code_block(doc, '''# Run today's A-share selection
cd "AMINQT CODES/skills/a-share-selection-strategy-skill"
python skills/a-share-selection-strategy/scripts/run_today_a_share_selection.py

# Run with specific symbols
python scripts/run_today_a_share_selection.py \\
    --history-source akshare \\
    --symbols 000001.SZ,000002.SZ,600519.SH

# Run with local prices file
python scripts/run_today_a_share_selection.py \\
    --prices-input prices.csv

# Score candidates only (no full scan)
python scripts/score_candidates.py \\
    --input prices.csv \\
    --mode generic

# Validate OHLCV data
python scripts/validate_ohlcv.py \\
    --input prices.csv''')

add_heading_styled(doc, 'Key Artifacts Produced', level=2)
add_bullet(doc, 'Execution manifest with full provenance', bold_prefix='run_manifest.json — ')
add_bullet(doc, 'Summary with execution path, coverage, failures', bold_prefix='summary.json — ')
add_bullet(doc, 'Final ranked candidate list', bold_prefix='candidates.csv — ')
add_bullet(doc, 'Per-symbol diagnostics and gate results', bold_prefix='diagnostics.csv — ')
add_bullet(doc, 'Human-readable HTML report', bold_prefix='report.html — ')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 7. Project Built-in Skills
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, '7. Project Built-in Skills', level=1)
add_para(doc, 'These skills live in the project root skills/ folder. The AI assistant loads them automatically based on task context.')

builtin_data = [
    ['Skill File', 'Trigger', 'What It Does'],
    ['quant-factor-engineering.md', 'Writing factor_engine.py', 'MACD/KDJ/BOLL/RSI specs, derivative features, sliding window'],
    ['quant-data-pipeline.md', 'Writing data_loader.py', 'akshare column mapping, incremental update, data validation'],
    ['risk-circuit-breaker.md', 'Writing risk_filter.py / executor', '3-layer risk control, circuit breaker, T+1, audit log'],
    ['verify-before-done.md', 'Before claiming "done"', 'Must run verification commands and show evidence'],
    ['graph-to-vector.md', 'K-line graph vectorization', '4 paths: matrix, CNN, patterns, graph embedding'],
]
btable = doc.add_table(rows=len(builtin_data), cols=3)
btable.style = 'Light Shading Accent 1'
for i, row_data in enumerate(builtin_data):
    for j, cell_text in enumerate(row_data):
        cell = btable.rows[i].cells[j]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

add_heading_styled(doc, 'How AI Loads These Skills', level=2)
add_code_block(doc, '''# The AI assistant reads skills/README.md at session start
# Then loads relevant skill based on the current task:

# Example: When you ask "implement MACD indicator"
# AI detects: factor engineering task
# AI loads: skills/quant-factor-engineering.md
# AI follows: the checklist in that file

# Example: When you ask "download stock data"
# AI detects: data pipeline task
# AI loads: skills/quant-data-pipeline.md
# AI follows: column mapping rules + validation checklist

# Example: When AI says "I'm done"
# AI must load: skills/verify-before-done.md
# AI must: run verification commands, show output, then claim done''')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 8. Quick Reference Card
# ════════════════════════════════════════════════════════════
add_heading_styled(doc, '8. Quick Reference Card', level=1)

add_heading_styled(doc, 'Python Libraries — One-Line Import', level=2)
add_code_block(doc, '''# quant-ohlcv-feature (335+ indicators)
from momentum_feature.Rsi import signal as rsi; df = rsi(df, 14, "rsi_14")

# ohlcpattern (40+ candlestick patterns)
from ohlcpattern.candlestick import CandlestickPatterns; csp = CandlestickPatterns(df)

# aurumq-rl (RL stock selection)
from aurumq_rl import FactorPanelLoader; panel = FactorPanelLoader.build_synthetic(60, 50, 12)

# quantitative_strategy_mcp (MCP tools)
from mcp_tools import call_mcp_tool; result = call_mcp_tool("get_stock_data", ts_code="000001.SZ")''')

add_heading_styled(doc, 'CLI Commands — One-Line Run', level=2)
add_code_block(doc, '''# aurumq-rl train
aurumq-rl-train --config configs/default.yaml

# aurumq-rl inference
aurumq-rl-infer --model-path models/best.onnx --data data/raw/

# aurumq-rl quickstart (CPU, 30s)
python "AMINQT CODES/skills/aurumq-rl/examples/quickstart.py"

# ohlcpattern CLI
ohlcpattern extract data/raw/600519.csv --output patterns.csv

# a-share-selection-strategy
python "AMINQT CODES/skills/a-share-selection-strategy-skill/skills/a-share-selection-strategy/scripts/run_today_a_share_selection.py"

# uzi-skill deep analysis
python "AMINQT CODES/skills/investagent/skills/uzi-skill/run.py" 600519.SH

# TradingAgents multi-agent
python "AMINQT CODES/skills/investagent/skills/TradingAgents/main.py" --ticker NVDA --date 2024-01-15''')

add_heading_styled(doc, 'AI Skill Triggers — Say These in Chat', level=2)
add_code_block(doc, '''# investagent (full-stack research)
"Research A-share AI semiconductor sector"
"Deep analysis of Kweichow Moutai"
"Buffett framework analysis of Apple"
"Multi-agent analysis of NVDA"
"Backtest a dual-MA strategy on BTC"

# a-share-selection-strategy
"Today A-share selection"
"Select low-price stocks"
"Expand stock pool scan"
"Local scoring with prices.csv"''')

# ── Save ──
output_path = os.path.join("AMINQT CODES", "skills", "skill.docx")
doc.save(output_path)
print(f"Saved to: {output_path}")
