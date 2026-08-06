# -*- coding: utf-8 -*-
"""_gen_feature_report.py — 生成《特征全貌调研报告》Word 文档 (WORM 日期后缀).

素材来源 (2026-08-04 实地核查):
  - feature_selector.py FREQ_ASSIGNMENT(69 实测) / FAMILY_ANALOG(类比)
  - _classify_freq_full.py 全市场×3年 6格 + 事件池结论 (memory 落档)
  - factor_registry 生产精选: main 1066 / dual 30 → {月/周/日/事件} 路由
  - 三个只读调查智能体: DUAL 事件根因 / 类比列盘点 / 模型存储架构

用法: python scripts/_gen_feature_report.py
输出: D:/AMINQT/REERENCE/DESIGN ALL/FEATURE/特征全貌调研报告_<ts>.docx
"""

import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)) + "/..")

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_DIR = r"D:/AMINQT/REFERENCE/Design All/FEATURE"


# ────────────────────────────── 数据素材 ──────────────────────────────

# 69 已实测列 (与 feature_selector.FREQ_ASSIGNMENT 逐字一致)
FREQ_ASSIGNMENT = [
    # (列, 频率, 类型, 族)
    ("pct_90_con", "月", "TS", "筹码"),
    ("pct_90_high", "月", "TS", "筹码"),
    ("weight_avg", "月", "TS", "筹码"),
    ("conc_trend_20d", "月", "TS", "筹码"),
    ("resistance_dist", "月", "TS", "筹码"),
    ("chip_entropy", "日", "TS", "筹码"),
    ("chip_gini", "日", "TS", "筹码"),
    ("peak_roc_5d", "日", "TS", "筹码"),
    ("chip_skew_dist", "月", "XS", "筹码"),
    ("conc_90_industry_rank", "月", "XS", "筹码"),
    ("peak_price", "月", "XS", "筹码"),
    ("peak_roc_20d", "周", "TS", "筹码"),
    ("cost_bias", "周", "XS", "筹码"),
    ("support_dist", "周", "TS", "筹码"),
    ("cost_50pct", "月", "TS", "成本线"),
    ("cost_95pct", "月", "TS", "成本线"),
    ("close_hfq", "周", "TS", "价格"),
    ("volume", "月", "XS", "量"),
    ("amount", "月", "XS", "量"),
    ("turnover_rate", "月", "XS", "量"),
    ("free_float_turnover_rate", "月", "XS", "量"),
    ("volume_ratio", "月", "TS", "量"),
    ("ma_vol_ratio_5_20", "月", "XS", "量"),
    ("vol_surge", "月", "XS", "量"),
    ("amt_surge", "月", "XS", "量"),
    ("bias_5", "周", "TS", "均线乖离"),
    ("bias_10", "周", "TS", "均线乖离"),
    ("bias_20", "月", "TS", "均线乖离"),
    ("bias_60", "日", "TS", "均线乖离"),
    ("bias_120", "月", "TS", "均线乖离"),
    ("bias_250", "月", "TS", "均线乖离"),
    ("bias_5_20_cross", "日", "XS", "均线乖离"),
    ("bias_20_60_cross", "日", "TS", "均线乖离"),
    ("amplitude_5d", "月", "XS", "波动"),
    ("intraday_range", "月", "XS", "波动"),
    ("winner_ratio", "周", "TS", "波动"),
    ("pctChg", "周", "TS", "波动"),
    ("pe_ttm", "周", "TS", "估值市值"),
    ("pb", "周", "TS", "估值市值"),
    ("ps_ttm", "周", "TS", "估值市值"),
    ("dv_ratio", "月", "TS", "估值市值"),
    ("total_mv", "周", "TS", "估值市值"),
    ("circ_mv", "周", "TS", "估值市值"),
    ("margin_balance", "月", "TS", "两融"),
    ("short_balance", "周", "TS", "两融"),
    ("margin_buy_amt", "月", "XS", "两融"),
    ("short_sell_vol", "周", "TS", "两融"),
    ("roe", "月", "TS", "基本面-盈利"),
    ("roe_deducted", "月", "TS", "基本面-盈利"),
    ("roa", "月", "TS", "基本面-盈利"),
    ("gross_margin", "月", "TS", "基本面-盈利"),
    ("debt_ratio", "月", "TS", "基本面-盈利"),
    ("current_ratio", "月", "TS", "基本面-盈利"),
    ("asset_turnover", "月", "TS", "基本面-盈利"),
    ("ar_turnover", "月", "TS", "基本面-盈利"),
    ("inventory_turnover", "月", "TS", "基本面-盈利"),
    ("rev_yoy", "周", "TS", "基本面-成长"),
    ("net_margin", "周", "TS", "基本面-成长"),
    ("eps_yoy", "周", "TS", "基本面-成长"),
    ("profit_yoy", "周", "TS", "基本面-成长"),
    ("ocfps", "周", "TS", "基本面-成长"),
    ("revenue_ps", "周", "TS", "基本面-成长"),
    ("bps", "周", "TS", "基本面-成长"),
    ("eps", "周", "TS", "基本面-成长"),
    ("dt_eps", "周", "TS", "基本面-成长"),
    ("roe_yoy", "周", "TS", "基本面-成长"),
    ("q_roe", "周", "TS", "基本面-成长"),
    ("q_ocf_to_sales", "周", "TS", "基本面-成长"),
    ("ocf_to_or", "日", "TS", "基本面-现金流"),
]

# 84 类比列 (与 feature_selector.FAMILY_ANALOG 一致) + Agent B 盘点
# (列, 频率, 类型, 族, 板块, 数据源/生成位置, 备注)
FAMILY_ANALOG = [
    # 价格族
    ("open", "周", "TS", "价格", "main", "持久化面板", "同 close_hfq"),
    ("high", "周", "TS", "价格", "main", "持久化面板", ""),
    ("low", "周", "TS", "价格", "main", "持久化面板", ""),
    ("close", "周", "TS", "价格", "main", "持久化面板", ""),
    ("pre_close", "周", "TS", "价格", "main", "持久化面板", ""),
    ("open_hfq", "周", "TS", "价格", "main", "持久化面板", ""),
    ("high_hfq", "周", "TS", "价格", "main", "持久化面板", ""),
    ("low_hfq", "周", "TS", "价格", "main", "持久化面板", ""),
    # 换手/量/流动性
    (
        "turn",
        "月",
        "XS",
        "换手",
        "main",
        "孤儿(scripts/_deactivate_orphan_autoadopted.py:23)",
        "无生成位置",
    ),
    (
        "turnover_rate_f",
        "月",
        "XS",
        "换手",
        "main",
        "data_supply.py:2190",
        "daily_basic",
    ),
    (
        "rank_ff_turnover",
        "月",
        "XS",
        "换手",
        "main",
        "cleaning_pipeline.py:122",
        "板块内 rank",
    ),
    (
        "rank_amount",
        "月",
        "XS",
        "换手",
        "main",
        "cleaning_pipeline.py:128",
        "板块内 rank",
    ),
    (
        "liquidity_score",
        "月",
        "XS",
        "换手",
        "main",
        "cleaning_pipeline.py:130",
        "板块内 rank",
    ),
    (
        "adv20",
        "月",
        "XS",
        "换手",
        "main",
        "feature_engine dim19:1786",
        "注释标『中间量非特征』",
    ),
    (
        "turnover_stability_5",
        "月",
        "XS",
        "换手",
        "main",
        "cleaning_pipeline.py:139",
        "",
    ),
    # 股本/市值
    ("total_share", "周", "TS", "股本", "main", "持久化面板", ""),
    ("float_share", "周", "TS", "股本", "main", "持久化面板", ""),
    ("free_share", "周", "TS", "股本", "main", "持久化面板", ""),
    # 筹码/成本 _x/_y (Tushare cyq merge 产物)
    (
        "pct_70_con_x",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "基列在 _XRANK_WHITELIST:702-716",
    ),
    (
        "pct_70_con_y",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "pct_70_high_x",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "pct_70_high_y",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "pct_70_low_x",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "pct_70_low_y",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "pct_90_con_x",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "pct_90_con_y",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "pct_90_high_x",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "pct_90_high_y",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "pct_90_low_x",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "pct_90_low_y",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    ("cost_5pct_x", "月", "TS", "筹码变体", "main", "dim21_chip_tushare:1797-1837", ""),
    ("cost_5pct_y", "月", "TS", "筹码变体", "main", "dim21_chip_tushare:1797-1837", ""),
    (
        "cost_15pct_x",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "cost_15pct_y",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "cost_50pct_x",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "cost_50pct_y",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "cost_85pct_x",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "cost_85pct_y",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "cost_95pct_x",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "cost_95pct_y",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    ("avg_cost_x", "月", "TS", "筹码变体", "main", "dim21_chip_tushare:1797-1837", ""),
    ("avg_cost_y", "月", "TS", "筹码变体", "main", "dim21_chip_tushare:1797-1837", ""),
    (
        "weight_avg_x",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    (
        "weight_avg_y",
        "月",
        "TS",
        "筹码变体",
        "main",
        "dim21_chip_tushare:1797-1837",
        "",
    ),
    # 股息 / 涨跌停
    ("dv_ttm", "月", "TS", "股息", "main", "持久化面板", "同 dv_ratio"),
    ("up_limit_raw", "日", "TS", "涨跌停", "main", "持久化面板", "快信号"),
    ("down_limit_raw", "日", "TS", "涨跌停", "main", "持久化面板", "快信号"),
    # 行业/市场 (窗口后缀定频)
    (
        "sw_ret_1d",
        "日",
        "TS",
        "行业",
        "dual",
        "panel_builder.py:523-525",
        "申万指数广播",
    ),
    (
        "sw_ret_1d_x",
        "日",
        "TS",
        "行业",
        "dual",
        "panel_builder.py:523-525",
        "merge 产物",
    ),
    (
        "sw_ret_5d",
        "周",
        "TS",
        "行业",
        "dual",
        "疑孤儿",
        "仅 eval 脚本引用; dim28 产 sw_lN_*",
    ),
    ("sector_return", "周", "TS", "行业", "dual", "dim04_sector_effect:1057-1059", ""),
    (
        "sector_return_5d",
        "周",
        "TS",
        "行业",
        "dual",
        "dim04_sector_effect:1057-1059",
        "",
    ),
    ("sw_relative_strength", "周", "TS", "行业", "dual", "疑孤儿", ""),
    ("sw_rotation_position", "周", "TS", "行业", "dual", "疑孤儿", ""),
    ("sw_ret_20d", "月", "TS", "行业", "dual", "疑孤儿", ""),
    ("sw_vol_20d", "月", "TS", "行业", "dual", "疑孤儿", ""),
    (
        "ind_holder_trend_20d",
        "月",
        "TS",
        "行业",
        "dual",
        "dim27_industry_flow:2333-2374",
        "",
    ),
    ("sw_momentum_accel", "月", "TS", "行业", "dual", "疑孤儿", ""),
    (
        "ind_margin_accel",
        "月",
        "TS",
        "行业",
        "dual",
        "dim27_industry_flow:2333-2374",
        "",
    ),
    (
        "ind_margin_chg_5d",
        "周",
        "TS",
        "行业",
        "dual",
        "dim27_industry_flow:2333-2374",
        "",
    ),
    ("sw_index_close", "周", "TS", "行业", "dual", "panel_builder.py:523-525", ""),
    ("sw_index_vol", "周", "TS", "行业", "dual", "panel_builder.py:523-525", ""),
    (
        "market_turnover",
        "周",
        "TS",
        "市场",
        "dual",
        "dim14_market_sentiment:1483-1489",
        "全市场聚合",
    ),
    (
        "market_turnover_ratio_5d",
        "周",
        "TS",
        "市场",
        "dual",
        "dim14_market_sentiment:1483-1489",
        "",
    ),
    (
        "market_turnover_ratio_20d",
        "月",
        "TS",
        "市场",
        "dual",
        "dim14_market_sentiment:1483-1489",
        "",
    ),
    (
        "market_limit_up",
        "日",
        "TS",
        "市场",
        "dual",
        "dim14_market_sentiment:1483-1489",
        "",
    ),
    # 波动/流动性
    (
        "ATR_pct",
        "月",
        "XS",
        "波动",
        "main",
        "dim02_volatility:982",
        "逐股 TR/close (非 dual!)",
    ),
    ("amihud_illiq", "月", "XS", "流动性", "main", "dim17:1655", ""),
    ("amihud_illiquidity", "月", "XS", "流动性", "main", "dim19:1783", ""),
    ("sw_turnover_anomaly", "月", "XS", "流动性", "dual", "疑孤儿", ""),
    (
        "free_float_turnover_rate_xrank",
        "月",
        "XS",
        "流动性",
        "main",
        "_add_cross_sectional_ranks:815",
        "date+board rank",
    ),
    (
        "amount_xrank",
        "月",
        "XS",
        "流动性",
        "main",
        "_add_cross_sectional_ranks:815",
        "",
    ),
    (
        "turnover_f_chg_5d",
        "月",
        "XS",
        "流动性",
        "main",
        "feature_engine dim05:1099",
        "",
    ),
    # 快价格信号
    ("close_vs_low", "日", "TS", "快信号", "main", "dim20:1732", ""),
    ("overnight_ret", "日", "TS", "快信号", "main", "dim20:1697", ""),
    ("ROC_3d", "日", "TS", "快信号", "main", "孤儿", "无生成位置"),
    ("gap_strength_5d", "周", "TS", "快信号", "main", "dim30:3124-3125", ""),
    ("gap_strength_20d", "月", "TS", "快信号", "main", "dim30:3124-3125", ""),
    ("gap_vs_ma5", "周", "TS", "快信号", "main", "dim20:1725", ""),
    # 其他
    ("month", "月", "TS", "日历", "dual", "dim08_calendar_month:1237", "季节效应"),
    (
        "list_days",
        "月",
        "TS",
        "上市天数",
        "main",
        "ingest_scan.py:46",
        "V3 已移除 (data_supply.py:2499)",
    ),
    (
        "benefit_part_x",
        "周",
        "TS",
        "赢家占比",
        "main",
        "cyq_ext.py:166/dim21:1815-1837",
        "= winner_ratio 旧名 (settings.py:44)",
    ),
    (
        "benefit_part_y",
        "周",
        "TS",
        "赢家占比",
        "main",
        "cyq_ext.py:166/dim21:1815-1837",
        "",
    ),
    (
        "churn_suspect",
        "月",
        "XS",
        "洗盘标志",
        "main",
        "cleaning_pipeline.py:140",
        "换手稳定性派生",
    ),
]

# 事件桶 (main 100 特征) — 按基列聚合的变体模式
EVENT_BUCKET = {
    "holder_count": [
        "pct10",
        "pct60",
        "std5",
        "std10",
        "std20",
        "std40",
        "max40",
        "d1",
        "d5",
        "d20",
    ],
    "sh_change_vol": [
        "pct2",
        "pct10",
        "pct60",
        "std5",
        "std10",
        "std40",
        "max40",
        "d1",
        "d5",
        "d20",
    ],
    "sh_change_amt": ["pct60", "std5", "std20", "std40", "max40", "d1", "d5", "d20"],
    "sh_net_change_sign": [
        "pct2",
        "pct5",
        "pct20",
        "pct60",
        "std5",
        "std10",
        "std40",
        "min40",
    ],
    "sh_change_amt_total": [
        "pct60",
        "std5",
        "std20",
        "std40",
        "max40",
        "d1",
        "d5",
        "d20",
    ],
    "sh_net_sign": [
        "pct2",
        "pct5",
        "pct20",
        "pct60",
        "std5",
        "std10",
        "std40",
        "mom20",
    ],
    "lhb_net_buy": [
        "ma5",
        "ma10",
        "ma20",
        "ma40",
        "ma60",
        "std20",
        "std40",
        "max10",
        "min10",
        "max20",
        "min20",
        "max40",
        "min40",
        "ema5",
        "ema20",
        "ema40",
    ],
    "lhb_buy_amt": [
        "ma5",
        "ma10",
        "ma20",
        "ma40",
        "ma60",
        "std20",
        "std40",
        "max10",
        "min10",
        "max20",
        "min20",
        "max40",
        "min40",
        "ema5",
        "ema20",
        "ema40",
    ],
    "lhb_sell_amt": [
        "ma5",
        "ma10",
        "ma20",
        "ma40",
        "ma60",
        "std20",
        "std40",
        "max10",
        "min10",
        "max20",
        "min20",
        "max40",
        "min40",
        "ema5",
        "ema20",
        "ema40",
    ],
}

# dual 30 特征 (selected_dual_20260730T115834.json)
DUAL_30 = [
    "market_turnover",
    "market_turnover_ratio_5d",
    "market_turnover_ratio_20d",
    "sw_ret_20d",
    "sw_momentum_accel",
    "ind_holder_trend_20d",
    "sw_vol_20d",
    "sw_ret_5d",
    "sector_return",
    "ATR_pct",
    "sw_ret_1d_x",
    "sw_turnover_anomaly",
    "sector_return_5d",
    "free_float_turnover_rate_xrank",
    "ind_margin_chg_5d",
    "ind_margin_accel",
    "amihud_illiq",
    "amount_xrank",
    "month",
    "market_limit_up",
    "sw_relative_strength",
    "sw_rotation_position",
    "gap_strength_20d",
    "gap_vs_ma5",
    "turnover_f_chg_5d",
    "amihud_illiquidity",
    "close_vs_low",
    "ROC_3d",
    "gap_strength_5d",
    "overnight_ret",
]


# ────────────────────────────── 文档助手 ──────────────────────────────


def set_cell_bg(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def add_table(doc, header, rows, widths=None, font_size=9, header_bg="D9E2F3"):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(font_size)
        set_cell_bg(hdr[i], header_bg)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(font_size)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def para(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p


# ────────────────────────────── 生成 ──────────────────────────────


def main() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()
    # 中文字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ── 封面 ──
    title = doc.add_heading("A股量化系统 · 特征全貌调研报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(
        doc,
        f"生成时间: {ts}    适用范围: 四模型架构 (月/周/日/事件) 宽表拆分前奏",
        size=9,
    )
    para(
        doc,
        "素材来源: feature_selector.py FREQ_ASSIGNMENT/FAMILY_ANALOG · "
        "_classify_freq_full.py 全市场×3年 6格+事件池 · factor_registry 生产精选 "
        "(main 1066 / dual 30) · 三个只读调查智能体 (DUAL事件/类比列盘点/模型存储)",
        size=9,
    )
    doc.add_paragraph()

    # ── 1. 背景与方法论 ──
    h1(doc, "1. 背景与判定方法论")
    para(
        doc,
        "系统正从『单频日频模型』(旧系统, 保存不动) 过渡到『四模型融合』(新系统: "
        "月频 + 周频 + 日频 + 事件四个平行子模型, DUAL 与 MAIN 合并进同一张宽表)。"
        "本报告盘点所有特征的事实状态, 作为 4 张频率宽表拆分的依据。",
    )
    h2(doc, "1.1 6格分类 (频率判定)")
    para(
        doc,
        "每特征在 6 格 {TS个股时序, XS日截面} × {日/周/月} 中取 |IC| 最强格; "
        "日/周/月用 1/5/20 天 pct_change 窗口 (label_pm_{2,5,20}d_net) 代理。"
        "TS = per-stock time-series IC (个股自身序列, 全样本平均); XS = 日截面 rank IC。",
    )
    h2(doc, "1.2 事件池 (事件类特征判定)")
    para(
        doc,
        "EVENT 特征 (LHB/HOLDER/BT) 不能用个股 TS 或日截面处理 —— 事件稀疏、序列常年为 0, "
        "per-stock TSIC 产生尖峰→反转结构性假象。正确做法 = 事件池时间对齐研究: "
        "聚合同类事件, 每事件取前1个月→后1个月窗口, 全部事件样本放一起找特征作用。",
    )

    # ── 2. 69 已实测列 ──
    h1(doc, "2. 已 6格实测列 (69 列, FREQ_ASSIGNMENT)")
    para(
        doc,
        "下列频率与类型均为 2026-08-04 全市场×3年真实判定 (scripts/_classify_freq_full.py), "
        "非猜测。brute-force 变体继承基列频率。",
    )
    rows = [(c, f, t, g) for c, f, t, g in FREQ_ASSIGNMENT]
    add_table(doc, ["特征列", "频率", "口径", "族"], rows, widths=[6.5, 2, 2, 4.5])
    para(doc, "")
    h2(doc, "2.1 组级判定汇总 (含 IC 数值)")
    para(
        doc,
        "月频模型 (TS·月): 筹码 weight_avg −0.044 / pct_90_con +0.024 / pct_90_high −0.029 / "
        "conc_trend_20d +0.015 / resistance_dist +0.015; 成本线 cost_50pct −0.038 / cost_95pct −0.032; "
        "均线 bias_250 −0.025 / bias_120 −0.013 / bias_20 −0.010; 股息 dv_ratio +0.046 (全表最强正); "
        "盈利质量 gross_margin +0.026 / asset_turnover +0.027 / roa +0.019 / roe +0.017; 两融 balance −0.032。",
    )
    para(
        doc,
        "周频模型 (TS·周): 价格 close_hfq −0.043; 估值负向 pe/pb/ps_ttm −0.042, "
        "total_mv/circ_mv −0.042; 成长 revenue_ps +0.030 / eps +0.016 / profit_yoy; "
        "winner_ratio −0.018; short_balance −0.022。",
    )
    para(
        doc,
        "日频模型 (TS·日弱): chip_entropy / chip_gini / peak_roc_5d; bias_5_20_cross (XS·日 +0.012); ocf_to_or。",
    )
    para(
        doc,
        "截面类负向 (XS·月, 归月频负向侧): 量全部负 amount −0.036 / volume −0.033 / turnover_rate −0.033 / "
        "vol_surge·amt_surge −0.022; 波动 amplitude_5d −0.026 / intraday_range −0.024; margin_buy_amt −0.029。",
    )

    # ── 3. 事件池结论 ──
    h1(doc, "3. 事件类特征 · 事件池结论")
    add_table(
        doc,
        ["事件组", "事件数", "事件定义 (铁律)", "事件池 IC 结论", "定位"],
        [
            (
                "LHB 龙虎榜",
                "23,691",
                "lhb_* 任一非空",
                "lhb_net_buy T+2 +0.156 → T+20 +0.045; top_buy +0.097 / top_sell −0.097",
                "正向 alpha (强)",
            ),
            (
                "HOLDER 增减持",
                "9,452",
                "仅稀疏 ratio 列 (sh_net_ratio/sh_g_ratio/sh_p_ratio/"
                "sh_c_ratio) 任一非空",
                "sh_c_ratio +0.048 / sh_net_sign +0.047 / sh_net_ratio +0.039, 一致正",
                "正向 alpha (月级)",
            ),
            (
                "BT 大宗",
                "25,584",
                "4 个 bt_ 列任一非空",
                "bt_disc_raw 一致负 −0.037; bt_amt_ratio_float_mv −0.022",
                "负向/风控信号",
            ),
        ],
        widths=[3.5, 2, 5.5, 5, 2.5],
    )
    para(doc, "")
    para(
        doc,
        "HOLDER 事件定义铁律: 事件只能用稀疏 ratio 列定义 (4 列 any = 9,452 事件, 覆盖率 0.41%); "
        "若含 sh_change_amt_total 等被 dim29 ffill 污染的列 → 1.3M 假事件 (56.9%), 严禁用于定义事件。"
        "被 ffill 污染的列仅在其真实事件池内评 IC。",
        bold=True,
    )

    # ── 4. FAMILY_ANALOG 类比列 ──
    h1(doc, "4. 同族类比映射列 (84 列, FAMILY_ANALOG)")
    para(
        doc,
        "下列频率为按同族函数一致性推导 (非单独 6格实测): 价格族→周 (同 close_hfq), "
        "换手/流动性→月·XS (同 turnover_rate), 股本→周 (同 total_mv), 筹码/成本 _x/_y→月 (同 pct_90_con), "
        "行业相对收益按窗口后缀定频 (_1d→日 / _5d→周 / _20d→月)。",
    )
    rows = [(c, f, t, g, b, s, n) for c, f, t, g, b, s, n in FAMILY_ANALOG]
    add_table(
        doc,
        ["特征列", "频率", "口径", "族", "板块", "数据源/生成位置", "备注"],
        rows,
        widths=[4.2, 1.4, 1.4, 2.2, 1.6, 5.2, 3.2],
        font_size=8,
    )
    para(doc, "")
    para(
        doc,
        "关键修正: ATR_pct 代码实为逐股 TR/close (dim02_volatility:982), 应归 main 非 dual; "
        "adv20 注释标『中间量非特征』; turn / ROC_3d / sw_ret_5d / sw_vol_20d / sw_relative_strength / "
        "sw_rotation_position / sw_momentum_accel / sw_turnover_anomaly 无生成位置 (疑孤儿), "
        "宽表构建时需过滤, 见 §7。",
    )

    # ── 5. 事件桶 100 ──
    h1(doc, "5. 事件桶 (main 精选 100 特征)")
    para(
        doc,
        "全部为事件基列的 brute-force 变体, 基列与变体窗口如下。事件表 MAIN+DUAL 合并时, "
        "同一批事件特征直接套用到双创股 (双创事件列由同一 feature_engine 生成)。",
    )
    rows = []
    total = 0
    for base, variants in EVENT_BUCKET.items():
        rows.append((base, ", ".join(variants), str(len(variants))))
        total += len(variants)
    add_table(doc, ["事件基列", "变体窗口", "数量"], rows, widths=[5, 10.5, 2])
    para(
        doc,
        f"合计 {total} 个事件特征。基列属: 增减持 (holder_count/sh_change_vol/sh_change_amt/"
        "sh_net_change_sign/sh_change_amt_total/sh_net_sign, 6 基列) + LHB (lhb_net_buy/"
        "lhb_buy_amt/lhb_sell_amt, 3 基列)。",
    )

    # ── 6. 生产精选路由 ──
    h1(doc, "6. 生产精选特征 · 四表路由结果")
    add_table(
        doc,
        ["板块", "月频表", "周频表", "日频表", "事件表", "未分类", "合计"],
        [("main", 535, 356, 75, 100, 0, 1066), ("dual", 15, 10, 5, 0, 0, 30)],
        widths=[3, 3, 3, 3, 3, 3, 3],
    )
    para(doc, "")
    para(
        doc,
        "main 1066 特征全为 brute-force 变体 (dedup 产出); dual 30 特征全为行业/市场相对类。"
        "合并后 4 张宽表: 月=550 / 周=366 / 日=80 / 事件=100 (去重叠前)。",
    )
    h2(doc, "6.1 dual 30 特征明细")
    para(doc, ", ".join(DUAL_30))
    para(doc, "")
    para(
        doc,
        "观察: 新版 MAIN 特征数偏高 (1066), 集中在月/周两张表; dual 偏少 (30)。"
        "对策见 §7 待办 (BruteForce 按频率限窗)。",
    )

    # ── 7. DUAL 事件桶根因 ──
    h1(doc, "7. DUAL 为什么没有事件桶 · 根因调查")
    para(doc, "结论: 不是双创缺数据, 是选择器杀列。", bold=True)
    para(
        doc,
        "① 生成层: feature_engine_v35.py:156-331 build() 无 board/股池条件, "
        "dim23/26/29/32/34 事件列对双创同样生成; cleaning_pipeline.py:83-95 step0_board_split "
        "按 symbol 前缀把 GEM+STAR 切为 dual, 事件列保留。",
    )
    para(
        doc,
        "② 杀列层: feature_selector.py:813-821 + _run_gate_d:936-951 —— dual 走 gate_d 管线, "
        "其 nan_filter(nan_threshold=0.95) 无 min_support 豁免 (行 436-458), 事件列 NaN≈98% 全被剔除; "
        "且 select() 快照在 apply_event_scope_screens 之前落盘 (train_runner.py:117-121), 事件筛选根本没机会生效。",
    )
    para(
        doc,
        "③ 双创事件数据量 (panel 266.7 万行): lhb_net_buy dual=6,338 / sh_net_ratio dual=5,705 / "
        "bt_count dual=14,704 非空, 均 > min_support=1000, 训练事件模型充足。",
    )
    para(
        doc,
        "④ 修复方向: gate_d 的 nan_filter 加 min_support 豁免 (复用 train_runner.py:86-91 MIN_SUPPORT=1000), "
        "或在四模型新路径 (QUAD-FREQ) 直接按事件基列路由建事件宽表, 不动旧 select()。",
    )

    # ── 8. 模型存储与命名 ──
    h1(doc, "8. 模型存储架构 · 两套模型并行命名")
    add_table(
        doc,
        ["系统", "目录", "文件名格式", "加载方式"],
        [
            (
                "UNI-FREQ 旧系统 (单频日频基线)",
                "models/pipeline1/ (零改动)",
                "{board}_{tag}.pkl 如 main_2026W31.pkl",
                "predict_runner.find_bundles: 按 board 前缀 + 字典序最大",
            ),
            (
                "QUAD-FREQ 新系统 (四频融合)",
                "models/pipeline1_quad/ (新建)",
                "{board}_quad_{freq}_{tag}.pkl (月/周/日/事件 4 子模型) "
                "+ {board}_quad_fused_{tag}.pkl (融合包)",
                "find_bundles/run_prediction 加 model_dir 参数",
            ),
        ],
        widths=[5, 3.5, 5, 4.5],
    )
    para(doc, "")
    para(
        doc,
        "当前磁盘: models/ 下仅 pipeline1/ 一个目录, 18 个 pkl (main_*/dual_*)。"
        "保存点 dual_track_trainer.py:637 DualTrackTrainer.save() → pickle bundle "
        "(board/feature_cols/models/calibrator/quantile/pain/rank); "
        "checkpoint 另存 .checkpoint_{board}_{tag}.pkl。"
        "DualTrackTrainer.__init__(model_dir=...) 已参数化, 新系统直接传 "
        "'models/pipeline1_quad' 即可; 旧路径与推理链路零改动。",
    )
    para(
        doc,
        "注意: config/settings.py:23 MODEL_DIR 指向 app/models/trained, 与 pipeline1 实际落盘不一致 (未使用), "
        "四频宽表/模型路径应以新目录为准并统一。",
    )

    # ── 9. 孤儿列与待办 ──
    h1(doc, "9. 孤儿列 / 风险清单 / 待办")
    h2(doc, "9.1 无生成位置 (宽表构建须过滤)")
    para(
        doc,
        "turn, ROC_3d, sw_ret_5d, sw_vol_20d, sw_relative_strength, sw_rotation_position, "
        "sw_momentum_accel, sw_turnover_anomaly — 但生产精选集含其中部分, 训练 df 上是否存在待核; "
        "adv20 为『中间量非特征』。",
    )
    h2(doc, "9.2 待办清单")
    para(doc, "① DUAL 事件桶修复 (nan_filter min_support 豁免) — 事件宽表覆盖双创。")
    para(
        doc,
        "② 扩展 6格分类实测 — 84 个类比列目前是同族推导, 用户要求真测替换猜测 (后台 ~40min)。",
    )
    para(
        doc,
        "③ BruteForce 按频率限窗 — 降 MAIN 特征数: 月列→20/40/60 窗, 周列→5/10/20, 日列→1/2/3/5。",
    )
    para(doc, "④ 4 张宽表构建 (月/周/日/事件, MAIN+DUAL 合并) + 4 子模型训练。")
    para(
        doc,
        "⑤ 融合方案 A/B/C 对比 — 检查指标 = 最终输出 STOCK TOP10 上涨幅度+概率准确度。",
    )
    para(doc, "⑥ 现有日频模型 (UNI-FREQ) 零改动保存; 新系统 QUAD-FREQ 并行目录。")

    # ── 附录 ──
    h1(doc, "附录 · 关键文件引用")
    for line in [
        "app/pipeline1/feature_selector.py — FREQ_ASSIGNMENT:665-705 / FAMILY_ANALOG:717-777 / _EVENT_PREFIXES:780 / freq_of",
        "app/pipeline1/feature_engine_v35.py — build():156-331, 事件 dims 23/26/29/32/34",
        "app/pipeline1/cleaning_pipeline.py — step0_board_split:83-95, 板块内 rank:122/128/130, churn_suspect:140",
        "app/pipeline1/train_runner.py — select_features:64-139 (NaN 预筛 MIN_SUPPORT:86-91, select 快照:117-121)",
        "app/pipeline1/dual_track_trainer.py — save():637, model_dir 默认 'models/pipeline1'",
        "app/pipeline1/predict_runner.py — find_bundles:41-74",
        "scripts/_classify_freq_full.py — GROUPS_EVT:72-95, _load:98-111 (全市场不按 board 切)",
        "factor_registry/selected_main_20260731T010208.json (1066) / selected_dual_20260730T115834.json (30)",
        "面板: D:/AMINQT/PARQUET/panel_full_enriched_v3.parquet (266.7万行 / 117列)",
        "模型: models/pipeline1/ (18 pkl, UNI-FREQ)",
    ]:
        para(doc, "· " + line, size=9)

    out = os.path.join(OUT_DIR, f"特征全貌调研报告_{ts}.docx")
    doc.save(out)
    print(f"已生成: {out}", flush=True)


if __name__ == "__main__":
    main()
