"""_gen_analog_summary.py — 类比列 6格分类实测总结 Word 文档 (WORM 日期后缀).

素材: scripts/_classify_freq_analog.py 实测结果 (2026-08-04):
  data/_classify_freq_analog_20260804_080206.log + _summary_<ts>.json
输出: D:/AMINQT/REFERENCE/Design All/Function Spec/FEATURE/类比列6格分类实测总结_<ts>.docx
用法: python scripts/_gen_analog_summary.py
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_DIR = r"D:/AMINQT/REFERENCE/Design All/Function Spec/FEATURE"
TITLE = "类比列 6格分类实测总结"

# ── 42 列实测 (feature, 族, 实测判定, |IC|, 同族猜测, 是否一致) ──
TESTED = [
    ("open", "价格", "TS·月", 0.0645, "周:TS", False),
    ("high", "价格", "TS·月", 0.0675, "周:TS", False),
    ("low", "价格", "TS·月", 0.0649, "周:TS", False),
    ("close", "价格", "TS·月", 0.0688, "周:TS", False),
    ("pre_close", "价格", "TS·月", 0.0624, "周:TS", False),
    ("open_hfq", "价格", "TS·月", 0.0642, "周:TS", False),
    ("high_hfq", "价格", "TS·月", 0.0672, "周:TS", False),
    ("low_hfq", "价格", "TS·月", 0.0647, "周:TS", False),
    ("rank_ff_turnover", "换手/流动", "TS·月", 0.0319, "月:XS", False),
    ("rank_amount", "换手/流动", "XS·月", 0.0433, "月:XS", True),
    ("liquidity_score", "换手/流动", "XS·月", 0.0380, "月:XS", True),
    ("adv20", "换手/流动", "XS·日", 0.0451, "月:XS", False),
    ("turnover_stability_5", "换手/流动", "TS·月", 0.0116, "月:XS", False),
    ("total_share", "股本", "XS·周", 0.0045, "周:TS", False),
    ("float_share", "股本", "TS·月", 0.0114, "周:TS", False),
    ("free_share", "股本", "XS·周", 0.0057, "周:TS", False),
    ("dv_ttm", "股息", "TS·月", 0.0645, "月:TS", True),
    ("up_limit_raw", "涨跌停", "TS·月", 0.0621, "日:TS", False),
    ("down_limit_raw", "涨跌停", "TS·月", 0.0620, "日:TS", False),
    ("sector_return", "行业/市场", "TS·周", 0.0092, "周:TS", True),
    ("sector_return_5d", "行业/市场", "TS·日", 0.0065, "周:TS", False),
    ("ind_holder_trend_20d", "行业/市场", "TS·日", None, "月:TS", False),
    ("ind_margin_accel", "行业/市场", "TS·日", 0.0125, "月:TS", False),
    ("ind_margin_chg_5d", "行业/市场", "TS·日", 0.0091, "周:TS", False),
    ("market_turnover", "行业/市场", "TS·月", 0.0308, "周:TS", False),
    ("market_turnover_ratio_5d", "行业/市场", "TS·日", 0.0158, "周:TS", False),
    ("market_turnover_ratio_20d", "行业/市场", "TS·月", 0.0254, "月:TS", True),
    ("market_limit_up", "行业/市场", "TS·日", 0.0091, "日:TS", True),
    ("ATR_pct", "波动", "TS·月", 0.0226, "月:XS", False),
    ("amihud_illiq", "流动性", "TS·月", 0.0382, "月:XS", False),
    ("amihud_illiquidity", "流动性", "TS·月", 0.0389, "月:XS", False),
    ("free_float_turnover_rate_xrank", "换手/流动", "XS·月", 0.0336, "月:XS", True),
    ("amount_xrank", "换手/流动", "XS·月", 0.0470, "月:XS", True),
    ("turnover_f_chg_5d", "换手/流动", "TS·日", None, "月:XS", False),
    ("close_vs_low", "价格", "XS·月", 0.0133, "日:TS", False),
    ("overnight_ret", "快价格信号", "TS·周", 0.0040, "日:TS", False),
    ("ROC_3d", "快价格信号", "TS·周", 0.0073, "日:TS", False),
    ("gap_strength_5d", "快价格信号", "TS·月", 0.0096, "周:TS", False),
    ("gap_strength_20d", "快价格信号", "TS·周", 0.0159, "月:TS", False),
    ("gap_vs_ma5", "快价格信号", "TS·日", 0.0023, "周:TS", False),
    ("month", "日历", "TS·周", 0.0463, "月:TS", False),
    ("churn_suspect", "赢家占比", "TS·月", 0.0297, "月:XS", False),
]

# ── 42 列缺失 (V3 已删列, 训练 df 中不存在) ──
MISSING = [
    "turn",
    "turnover_rate_f",
    "pct_70_con_x",
    "pct_70_con_y",
    "pct_70_high_x",
    "pct_70_high_y",
    "pct_70_low_x",
    "pct_70_low_y",
    "pct_90_con_x",
    "pct_90_con_y",
    "pct_90_high_x",
    "pct_90_high_y",
    "pct_90_low_x",
    "pct_90_low_y",
    "cost_5pct_x",
    "cost_5pct_y",
    "cost_15pct_x",
    "cost_15pct_y",
    "cost_50pct_x",
    "cost_50pct_y",
    "cost_85pct_x",
    "cost_85pct_y",
    "cost_95pct_x",
    "cost_95pct_y",
    "avg_cost_x",
    "avg_cost_y",
    "weight_avg_x",
    "weight_avg_y",
    "sw_ret_1d",
    "sw_ret_1d_x",
    "sw_ret_5d",
    "sw_relative_strength",
    "sw_rotation_position",
    "sw_ret_20d",
    "sw_vol_20d",
    "sw_momentum_accel",
    "sw_index_close",
    "sw_index_vol",
    "sw_turnover_anomaly",
    "list_days",
    "benefit_part_x",
    "benefit_part_y",
]

# ── 34 处不一致里的关键修正 (族级) ──
KEY_FIXES = [
    (
        "价格族 8 列 (open/high/low/close/pre_close + hfq 变体)",
        "周:TS",
        "TS·月",
        "|IC| 0.062~0.069, 全表最强档, 与 close_hfq 同为价格但频率判定不同",
    ),
    ("涨跌停 up_limit_raw / down_limit_raw", "日:TS", "TS·月", "|IC| 0.0621 / 0.0620"),
    (
        "amihud_illiq / amihud_illiquidity",
        "月:XS",
        "TS·月",
        "|IC| 0.0382 / 0.0389 — 非横截面, 是个股时序",
    ),
    ("ATR_pct", "月:XS", "TS·月", "|IC| 0.0226"),
    ("market_turnover", "周:TS", "TS·月", "|IC| 0.0308"),
    ("churn_suspect", "月:XS", "TS·月", "|IC| 0.0297 — 非 XS"),
    ("month", "月:TS", "TS·周", "|IC| 0.0463"),
]

# ── 弱信号 / 不可靠判定 (不能当可靠频率依据) ──
WEAK = [
    ("ind_holder_trend_20d", "TS·日", None, "截面样本不足, IC 为 NaN"),
    ("turnover_f_chg_5d", "TS·日", None, "截面样本不足, IC 为 NaN"),
    ("total_share", "XS·周", 0.0045, "|IC| 过小"),
    ("free_share", "XS·周", 0.0057, "|IC| 过小"),
    ("overnight_ret", "TS·周", 0.0040, "近零"),
    ("ROC_3d", "TS·周", 0.0073, "近零"),
    ("gap_vs_ma5", "TS·日", 0.0023, "近零"),
]

# ── 8 个与猜测一致列 ──
MATCHED = [
    ("rank_amount", "XS·月"),
    ("liquidity_score", "XS·月"),
    ("dv_ttm", "TS·月"),
    ("sector_return", "TS·周"),
    ("market_turnover_ratio_20d", "TS·月"),
    ("market_limit_up", "TS·日"),
    ("free_float_turnover_rate_xrank", "XS·月"),
    ("amount_xrank", "XS·月"),
]


# ────────────── docx 辅助 ──────────────
def _set_cn_font(run, size=10.5, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Calibri"
    r = run._element.rPr
    rFonts = r.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        r.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _heading(doc, text, lvl):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_cn_font(run, size={1: 16, 2: 13, 3: 11.5}[lvl], bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def _para(doc, text, size=10.5, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_cn_font(run, size=size, bold=bold)
    run.font.italic = italic
    return p


def _table(doc, headers, rows, widths=None, cell_fmt=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        _set_cn_font(c.paragraphs[0].add_run(h), size=9.5, bold=True)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ""
            text = "—" if v is None else str(v)
            run = cells[j].paragraphs[0].add_run(text)
            _set_cn_font(run, size=9.5)
            if cell_fmt:
                cell_fmt(i, j, run, v)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    return t


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")

    doc = Document()
    doc.add_heading(TITLE, level=0)
    _para(
        doc,
        f"生成时间: {datetime.now():%Y-%m-%d %H:%M:%S}  ·  方法: TS/XS × 日/周/月 "
        f"(1/5/20 变化窗口) 取 |IC| 最强格  ·  数据: 全市场×3年训练序列",
        italic=True,
    )

    # 1. 任务与背景
    _heading(doc, "1. 任务与背景", 1)
    _para(
        doc,
        "feature_selector.FAMILY_ANALOG 共 84 个类比列, 此前频率/类型归属是按同族猜测"
        " (例如价格族跟 close_hfq 判周)。本次对全部 84 列跑真实 6格分类, 替换猜测, "
        "并与猜测对比得出 一致/不一致 清单。不一致列是升级 FAMILY_ANALOG → "
        "FREQ_ASSIGNMENT 的依据。",
        size=10,
    )

    # 2. 方法
    _heading(doc, "2. 方法口径", 1)
    _para(
        doc,
        "用真实训练序列 (CleaningPipeline.run_train → FeatureEngineV35.build → "
        "prepare_board_frame, main/dual 逐板块) 构建含标签的训练 df, 取最近 3 年。"
        "每列计算 1/5/20 日变化窗口, 分别做个股时序 rank IC (TS) 与日截面 rank IC (XS),"
        "加权 wIC = 0.45×IC2d + 0.35×IC3d + 0.2×IC5d, 取 |IC| 最强格为判定。"
        "构建按列裁剪、逐板块释放, 不做 brute-force generate, 无 OOM。",
        size=10,
    )

    # 3. 结果总览
    _heading(doc, "3. 结果总览", 1)
    _table(
        doc,
        ["项", "数量", "说明"],
        [
            ["实测列", 42, "训练 df 中存在, 拿到真实 6格判定"],
            ["其中 与猜测一致", 8, "判定与同族猜测吻合"],
            ["其中 不一致", 34, "判定与同族猜测不符, 需升级"],
            ["缺失列", 42, "训练 df 中不存在 — V3 已删列 (见 §6)"],
            ["合计", 84, "FAMILY_ANALOG 全量"],
        ],
        widths=[4, 2, 9],
    )

    # 4. 实测判定明细
    _heading(doc, "4. 实测判定明细 (42 列)", 1)

    def fmt_run(i, j, run, v):
        if j == 5:
            run.font.color.rgb = None
        if j == 2 and TESTED[i][5]:
            run.font.bold = True
        if j == 4 and TESTED[i][5]:
            run.font.bold = True

    _table(
        doc,
        ["特征", "族", "实测判定", "|IC|", "同族猜测", "一致"],
        [
            (f, fam, verdict, f"{ic:.4f}" if ic else "NaN", guess, "✅" if ok else "❌")
            for f, fam, verdict, ic, guess, ok in TESTED
        ],
        widths=[5.5, 2.6, 2.2, 2.0, 2.4, 1.4],
        cell_fmt=fmt_run,
    )

    # 5. 关键修正
    _heading(doc, "5. 34 处不一致中的关键修正", 1)
    _table(
        doc,
        ["特征族", "原猜测", "实测", "说明"],
        [(f, old, new, note) for f, old, new, note in KEY_FIXES],
        widths=[6, 2.4, 2.4, 5.2],
    )
    _para(
        doc,
        "一致列 (8): " + "、".join(f"{c} ({v})" for c, v in MATCHED) + "。",
        size=10,
        italic=True,
    )

    # 6. V3 已删列发现
    _heading(doc, "6. 42 个缺失列 = V3 已删列 (关键发现)", 1)
    _para(doc, "42 个缺失列在当前 V3 面板与构建中都不再生成, 不是构建 bug:", size=10)
    _para(
        doc,
        "· chip/cost/avg_cost/weight_avg 的 _x/_y 后缀列 = 旧 cyq(Tushare 筹码) merge 产物,"
        " dim21_chip_tushare 于 2026-08-02 V3 删列后不再产出;",
        size=10,
    )
    _para(
        doc,
        "· sw_ret_*/sw_index_*/sw_relative_strength/sw_rotation_position/"
        "sw_momentum_accel/sw_turnover_anomaly = 已删除的行业/市场列;",
        size=10,
    )
    _para(
        doc,
        "· list_days = V3 入库门移除 (近 150 天次新不进 V3);  turn/turnover_rate_f/"
        "benefit_part_* = 无生成位置或旧名。",
        size=10,
    )
    _para(
        doc,
        "根因: 生产精选文件 (factor_registry, 2026-07-31) 是 V3 删列 (2026-08-02) "
        "之前的过时快照, 仍引用这 42 个死列。",
        bold=True,
        size=10,
    )
    _para(
        doc,
        "结论: 新频率宽表构建必须剔除这 42 列; 下次重训会自动丢弃。同时 "
        "tests/test_freq_assignment.py 的 PRODUCTION_BASES 含 turn/benefit_part_x/"
        "sw_ret_1d 等死列, 清理 FAMILY_ANALOG 时需同步修改。",
        size=10,
    )
    _para(doc, "缺失列清单:", bold=True, size=10)
    _para(doc, "、".join(MISSING), size=9.5)

    # 7. 弱信号警告
    _heading(doc, "7. 弱信号 / 不可靠判定 (不能当频率依据)", 1)
    _table(
        doc,
        ["特征", "最强格", "|IC|", "原因"],
        [(f, v, f"{ic:.4f}" if ic else "NaN", note) for f, v, ic, note in WEAK],
        widths=[5.5, 2.4, 2.0, 6.1],
    )

    # 8. 架构转向
    _heading(doc, "8. 架构转向 (2026-08-04)", 1)
    _para(
        doc,
        "用户决定不做月度/周度独立模型。推荐方案: 日频模型吃全部频率特征 "
        "(长窗口特征作输入) + 事件模型独立 (LHB/HOLDER 事件池 alpha, BT 折价率负向风控),"
        "两路融合; 用最终输出 STOCK TOP10 的上涨幅度+概率准确度裁决融合方案。"
        "现有日频模型 (UNI-FREQ, models/pipeline1/) 保存不动, 新系统另存 (QUAD-FREQ)。",
        size=10,
    )

    # 9. 下一步
    _heading(doc, "9. 下一步", 1)
    _para(
        doc,
        "1) 把 34 个实测判定升级进 FREQ_ASSIGNMENT (替换 FAMILY_ANALOG 猜测);",
        size=10,
    )
    _para(
        doc,
        "2) 42 个 V3 死列从宽表/路由剔除, 同步清理 FAMILY_ANALOG 与测试死列;",
        size=10,
    )
    _para(doc, "3) 建频率宽表, 训日频(吃全频) + 事件两模型, TOP10 裁决融合。", size=10)

    # 10. 落盘索引
    _heading(doc, "10. 关联文件", 1)
    _table(
        doc,
        ["文件", "说明"],
        [
            ["scripts/_classify_freq_analog.py", "6格实测脚本 (可复跑)"],
            ["data/_classify_freq_analog_20260804_080206.log", "实测结果日志 (WORM)"],
            ["data/_classify_freq_analog_summary_<ts>.json", "结构化摘要 (WORM)"],
            ["scripts/_gen_feature_report.py", "特征全貌调研报告生成器"],
            ["特征全貌调研报告_20260804_070521.docx", "全貌调研报告 (同目录)"],
        ],
        widths=[8, 7],
    )

    out = os.path.join(OUT_DIR, f"{TITLE}_{ts}.docx")
    doc.save(out)
    print(f"落盘: {out}")
    print(f"段落 {len(doc.paragraphs)}, 表格 {len(doc.tables)}")


if __name__ == "__main__":
    main()
