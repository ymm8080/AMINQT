"""_shortlist_t5_t10.py — 今日最终短名单 (每板块 TOP-5) 前瞻预测输出.

数据源 = 最后一次 FULL RUN 回测产物 (用户 2026-08-05: "已有 full run 基于昨日数据
做出预测, 模块特征/一切都要保留, 不得丢失"). 本脚本**不重算特征、不重新选股**,
只读 FULL RUN 落盘产物, 给每只短名单股计算**前瞻预测** (最新 score 经 OOS 校准).

  - shortlist_main.csv / shortlist_dual.csv : 权威选股结果 (score/co_occur/rk/cut/est_wr)
  - stocks_{board}_{system}_oos.csv          : OOS 逐股 评分 + 已实现前视 MFE
    → 按 score 分位桶校准 → 每股逐视界 预期涨幅(MFE) + 达到概率 (前瞻, 非历史回看)
  - backtest.json                           : 系统级 OOS 逐视界 胜率/期望 (SUMMARY 段)

**概率口径 (2026-08-06 用户定案):** 概率 = 该股**逐股自然概率** (每股唯一真值)
= P(该股该视界已实现 MFE ≥ 固定绝对目标) via 逐股 Platt 平滑校准 — 替代旧桶级共享值
(同 score 桶内每股同值, 用户: "natural, not bulk probability"). 不是 P(MFE>0) —
那曾是 85-99% 被用户否决 "incorrectly high".

**制度自适应门 (2026-08-05 用户):** 输出哪个 (板块, 系统) 组合**不写死**, 用最新 OOS
市场数据 (固定视界净收益 label_pm) 每日判定: top-quantile 选股平均净收益 > 池基线 + margin
才保留; **判定只看主视界 (T+3)** — 2026-08-05 用户否决用 T+5/T+10 兜底 ("如果没有优势就
不要入选") → 主视界无优势组合则空仓观望, 不输出清单. 见 config REGIME_GATE.

顶部 SUMMARY = 主板+双创合并的单一集成排名清单 (rank/命中模块/预期涨幅/达到概率), 直接决策;
明细表格保留每板块 T-5 逐视界.

流程 (用户): 先给每只股 预期涨幅+达到概率 (预测), 再排名 — 绝不先排名后预测.
**排名键 (2026-08-07 定案):** 每股 pred_mag_10d — 共享 calibrate_mag10d
(score→label_pm_10d_net, T+10 close-to-close 校准幅度) 全板块日截面降序 → 每板块
TOP-5 (2026-08-14 收紧: 250d OOS 幅度前沿 top5 幅度/命中率/大涨率全优于 top10,
见 _diag_mag_frontier), 与并行 build_merged_shortlist 同源 (d10 c2c 定案). 候选池=全板块
(latest 截面全部有分股, 不再 TOP-30 预筛); 旧 top-N (狙击 top5 / 融合 top10, 按特征分)
仅作 systems/co_occur 元数据标注; score_w 降为平局裁决/展示.

**主视界 (2026-08-05 用户定案):** T+3 (短持 3 天). 排名权重 3d=0.40 最高.
入选门 = **纯 T+3 门** (2026-08-09 删 2d 视界后联合门退化为纯 T+3; 见 config
SHORTLIST_SCORE.select_gate): 保留 ⇔ T+3 预期涨幅 > t3_min.
平局裁决用 T+3. 三视界 T+3/5/10 预期涨幅+达到概率仍全部展示 (T+10 降为参考视界).

用法: python scripts/_shortlist_t5_t10.py [YYYYMMDD 选股日, 默认=full run 短名单日期]
可选: 传 symbol 列表 (空格分隔) 查看这些股今日预测 (未入选股模型今日无打分, 无预测).
输出: STOCK_LIST_DIR/parallel_shortlist_<date>__<module>.csv + STOCK LIST <date>__<module>.docx/.xlsx
      (module = current_meta.json 模型版本 tag, 供回归按模块分组评估)
"""

import os
import re
import sys
from pathlib import Path

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))

import numpy as np
import pandas as pd
import pyarrow.parquet as pq
from sklearn.linear_model import LinearRegression, LogisticRegression

from app.pipeline1.label_engine import COST, slippage_tier
from app.pipeline_parallel.calibration import calibrate_mag10d
from app.pipeline_parallel.config import FUSION, HORIZONS, SNIPER
from app.pipeline_parallel.scoring import pool_score
from config.settings import (
    DATA_DIR,
    DATA_OTHERS_DIR,
    REGIME_GATE,
    SHORTLIST_SCORE,
    STOCK_LIST_DIR,
)

try:
    from docx import Document
    from docx.shared import Pt
except Exception:  # python-docx 未安装时跳过 docx 输出
    Document = None

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter
except Exception:  # openpyxl 未安装时跳过 xlsx 输出
    Workbook = None

# FULL RUN 目录: 默认自动解析最新一次并行回测 (含 shortlist_main.csv 的最晚 ts 子目录),
# 供逐日连续运行 (EMA 平滑需各日 raw 历史); 找不到才回退硬编码默认.
_FULLRUN_HARD = DATA_OTHERS_DIR / "BACKTESTING RESULT" / "20260806_144240"


def _latest_fullrun_dir(prefix: str | None = None) -> Path | None:
    """最新并行 run_dir.

    prefix: 按日期前缀限定 (如 "20260813" → 只在该交易日 run_dir 里挑最新),
    自动化集成用它确保交付"当日"并行清单; 找不到该日期 → None (调用方大声失败).
    带 prefix 时不要求 shortlist_main.csv — 某板当日可能被拒/为空, 交付端负责标注原因.
    无 prefix → 历史最新且含 shortlist_main.csv (standalone 用法不变).
    """
    base = DATA_OTHERS_DIR / "BACKTESTING RESULT"
    pattern = f"{prefix}_*/" if prefix else "*/"
    try:
        cands = sorted(base.glob(pattern), key=lambda p: p.name, reverse=True)
    except OSError:
        return _FULLRUN_HARD if prefix is None else None
    if prefix is None:
        cands = [p for p in cands if (p / "shortlist_main.csv").exists()]
    if cands:
        return cands[0]
    return _FULLRUN_HARD if prefix is None else None


def _reject_reason(run_dir: Path, board: str) -> str:
    """某板当日短名单为空 (未接受/被退回) 的原因, 从 run_dir 证据推导.

    有该板 OOS 落盘 → 板块有候选/模型跑过, 但选股+排名门当日无入选 → 选股门原因;
    连 OOS 落盘都没有 → 板块当日无候选 (池为空/未参与并行选股).
    """
    has_eval = (run_dir / f"stocks_{board}_sniper_oos.csv").exists() or (
        run_dir / f"stocks_{board}_fusion_oos.csv"
    ).exists()
    if has_eval:
        return "候选未通过选股/排名门, 当日短名单为空"
    return "该板块当日无候选 (池为空/未参与并行选股)"


FULLRUN_DIR = _latest_fullrun_dir()
print(f"[fullrun] 使用 {FULLRUN_DIR}", flush=True)
BOARD_LABEL = {"main": "主板", "dual": "双创(创业板+科创板)"}
# 集成清单逐视界展示顺序 (时间序; HORIZONS = ("3d","5d","10d"))
HORIZON_ORDER = ("3d", "5d", "10d")
# OOS score→前视涨幅校准 (前瞻预测): 每股独立连续校准 (LinearRegression score→mfe),
# 替代 CAL_BINS 桶级共享值 (2026-08-06 用户: 价格预测必须每股独立, 非 bulk);
# CAL_MIN_N 保留用于 load_system_stats/regime_gate 最少样本门.
CAL_MIN_N = 5
# [2026-08-06] 达到概率改逐股自然概率 (用户: "natural, not bulk probability"):
# 口径 = P(该股该视界达到固定绝对涨幅) — 每股唯一真值 (score 单调略增, 非桶内共享).
ABS_TARGET = {"3d": 0.03, "5d": 0.04, "10d": 0.06}
# [2026-08-06] 价格预测=每股独立, 只用该股自己最近 ~6 个月历史 (用户定案):
# 每股取自己最近 PER_STOCK_WINDOW 交易日 (score→mfe) 拟合, 不足 PER_STOCK_MIN_N 回退横截面.
PER_STOCK_WINDOW = 130
PER_STOCK_MIN_N = 30
# [2026-08-06] 预测稳定性 (诊断 _diag_pred_decomp: 校准器逐日漂移 + score 逐日抖动同量级):
#  1) per-stock 斜率向横截面收缩 (empirical-Bayes partial pooling): λ=n/(n+SHRINK_KAPPA),
#     保持个股质心 (intercept=ȳ−slope·x̄) → 面板逐日滚动重拟合时 Δslope 大幅降低.
#  2) 输出级时间平滑 (EMA): 同日股用近 SMOOTH_K 个可用交易日 raw 预测的衰减加权均值
#     w_k=α·(1-α)^k (k=0=今日), gap-robust (缺日跳过) → 直接抑制相邻交易日预测/概率剧变.
SHRINK_KAPPA = 40
SMOOTH_ENABLED = True
SMOOTH_ALPHA = 0.35
SMOOTH_K = 12
# [2026-08-07] 短名单排名定案: 按每股 pred_mag_10d (共享 calibrate_mag10d, score→
# label_pm_10d_net, T+10 close-to-close 校准幅度) 排名, 替代原 score_w 混合排名.
# 250d OOS (scripts/_diag_parallel_rank_compare.py + _diag_rank_strategies.py) 证明:
# 纯 pred_mag 排名 MFE 全视界赢特征排名 +3.9~10.1pp, 上涨率打平; 混合/两段式全不如纯.
# 原短名单仅 ~15 只 (特征 TOP-5∪TOP-10), pred_mag 在内重排近乎无效 → 候选池加宽到
# 全板块 (latest 截面全部有分股), 再按 pred_mag_10d 取每板块 TOP-10 (TOP-30 预筛已废).
CAND_RANK_KEY = "pred_mag_10d"  # 排名键 = 每股 mag_10d 预期幅度 (2026-08-07 定案)
# [2026-08-14] 报告幅度锚定 (用户: 并行预测高于实际, 主板+双创都要修): pred_ret_{h}/
# pred_mag_10d 由短窗(cal_n=21)横截面 OLS 对今日热 score 外推 → 顺周期高估 (热行情单日
# 报告 10d 飙到 12-13%, 2-3x 于长期实得). 拉长校准窗 125d 证伪 (诊断 _diag_cal_window_compare:
# 近期 120d dual bias +2.84pp/实得 +0.14% vs cal_n=21 +0.23pp/+3.46%). 裁决: 排名键保持
# cal_n=21 (排序有真本事), 报告幅度平移至模型近 ANCHOR_WINDOW 决策日 top-ANCHOR_TOP
# 已实现均值. 锚深 = 入选深 (ANCHOR_TOP=5, 2026-08-14 与 Top-5 收紧同步: 报告的是用户
# 真正下单的 top-5 档位实得). 窗长 2026-08-14 120→250: 用户"3% 低估"质疑成立 — 250d
# 幅度前沿 (诊断 _diag_mag_frontier) 双创 top-5 实得 +6.34%/top-10 +5.76%, 主板 top-5
# +3.49%; 120d 恰踩双创弱行情段低估 (top-10 +3.37% vs 250d +5.76%). 250d 与项目验收
# 口径一致 (oos-only-acceptance), 报真实长期能力而非当下弱市; 热日尖峰已由锚的"平移至
# 实得"机制消除, 不会回到 12-13%.
ANCHOR_WINDOW = 250
ANCHOR_TOP = 5  # 锚定深度 = 入选深度 (每板块 TOP-5)


def load_system_stats(records: dict) -> dict:
    """OOS 逐股记录 → {(board, system): {h: {"mag", "hit", "n"}}} (SUMMARY 段, 系统级汇总).

    系统级单一值 (非个股值): mag=系统期望涨幅(均值), hit=P(已实现 MFE ≥ 系统期望) —
    不是 P(MFE>0) 的虚高系统胜率 (旧 82-94% 被用户否决 "not convincing")."""
    stats = {}
    for (board, key), rec in records.items():
        if key == "both":
            continue
        per = {}
        for h in HORIZONS:
            v = rec[f"mfe_{h}"].dropna()
            if len(v) < CAL_MIN_N:
                continue
            exp = float(v.mean())
            per[h] = {"mag": exp, "hit": float((v >= exp).mean()), "n": int(len(v))}
        stats[(board, key)] = per
    return stats


# 面板固定视界净收益列 (现实"持有到收盘"口径, 制度门用): label_pm_{2,3,5,10}d_net → pm_{h}
PM_LABEL = {h: f"label_pm_{h[:-1]}d_net" for h in HORIZONS}
PM_RENAME = {v: f"pm_{h}" for h, v in PM_LABEL.items()}


def _norm_oos(df: pd.DataFrame) -> pd.DataFrame:
    """OOS 逐股记录 → 留 symbol/date + score + 各视界已实现 MFE (前视未实现行留 NaN)."""
    d = df.rename(columns={f"label_mfe_{h}_net": f"mfe_{h}" for h in HORIZONS})
    keep = ["symbol", "date", "score"] + [f"mfe_{h}" for h in HORIZONS]
    return d[keep]


def _load_panel_pm(board: str) -> pd.DataFrame:
    """投影 3y 诊断面板的 label_pm 列 (symbol+date+四视界净收益), 供制度门并入 OOS 记录."""
    fp = DATA_DIR / f"_diag_stage_{board}_3y.parquet"
    t = pq.read_table(str(fp), columns=["symbol", "date"] + list(PM_LABEL.values()))
    t = t.to_pandas().rename(columns=PM_RENAME)
    t["symbol"] = t["symbol"].astype(str)
    t["date"] = pd.to_datetime(t["date"]).dt.strftime("%Y-%m-%d")
    return t


def load_oos_records() -> dict[tuple[str, str], pd.DataFrame]:
    """FULL RUN 的 OOS 逐股 评分 + 已实现前视 MFE + 面板固定视界净收益 label_pm
    → {(board, key): DataFrame}. key: "sniper" / "fusion" / "both"(共现=两系统合并)."""
    records = {}
    for board in ("main", "dual"):
        pm = _load_panel_pm(board)
        sn = _norm_oos(
            pd.read_csv(
                FULLRUN_DIR / f"stocks_{board}_sniper_oos.csv", dtype={"symbol": str}
            )
        )
        fu = _norm_oos(
            pd.read_csv(
                FULLRUN_DIR / f"stocks_{board}_fusion_oos.csv", dtype={"symbol": str}
            )
        )
        records[(board, "sniper")] = sn.merge(pm, on=["symbol", "date"], how="left")
        records[(board, "fusion")] = fu.merge(pm, on=["symbol", "date"], how="left")
        records[(board, "both")] = pd.concat(
            [records[(board, "sniper")], records[(board, "fusion")]], ignore_index=True
        )
    return records


def regime_gate(records: dict) -> dict:
    """制度自适应门: 用最新 OOS 市场数据判定每个 (board, system) 组合今日是否保留.

    每组合: base_ev_h = 池内全部行平均 label_pm (闭眼全买池基准);
    top_ev_h = 高分 top_quantile 选股平均 label_pm; 保留 ⇔ top_ev > base_ev + margin
    **只看主视界** (REGIME_GATE.primary_horizon) — 不许长视界兜底 (2026-08-05 用户:
    "如果没有优势就不要入选"). 全视界 alpha 仍进 per 供 SUMMARY 展示. 返回
    {(board, sys): {"active": bool, "per": {h: {...}}}}.
    """
    cfg = REGIME_GATE
    q, margin = cfg["top_quantile"], cfg["margin"]
    primary = cfg["primary_horizon"]
    out = {}
    for board in ("main", "dual"):
        for sname in ("sniper", "fusion"):
            rec = records.get((board, sname))
            if rec is None or len(rec) == 0:
                continue
            s = rec["score"]
            thr = s.quantile(1.0 - q)  # top q-quantile (高分选股)
            top_mask = s >= thr
            per = {}
            for h in cfg["horizons"]:
                col = f"pm_{h}"
                if col not in rec:
                    continue
                v = rec[col].dropna()
                if len(v) < CAL_MIN_N:
                    continue
                tv = rec.loc[top_mask, col].dropna()
                if len(tv) < CAL_MIN_N:
                    continue
                base_ev, top_ev = float(v.mean()), float(tv.mean())
                per[h] = {
                    "base_ev": base_ev,
                    "top_ev": top_ev,
                    "top_wr": float((tv > 0).mean()),
                    "alpha": top_ev - base_ev,
                }
            active = primary in per and per[primary]["alpha"] > margin
            out[(board, sname)] = {"active": active, "per": per}
    return out


def _row_active(row, gate: dict) -> bool:
    """该短名单行 (board, systems) 是否命中今日保留的组合; 共现股=任一所涉系统保留即算.

    2026-08-07: 全板块候选 systems="" (旧 top-N 之外) 走合并 "both" 校准 (score=max
    两系统池分) → 同 "both": 任一系统保留即算, 不因无系统标签被整批误杀.
    2026-08-09: 共现股 systems="fusion+sniper" 同样算任一系统保留 (原漏判 → 共现股被
    误标 未过, 已修)."""
    sys = str(row["systems"])
    combos = ("sniper", "fusion") if sys in ("", "both", "fusion+sniper") else (sys,)
    return any(gate.get((row["board"], s), {}).get("active", False) for s in combos)


def gate_fallback(gate: dict) -> tuple | None:
    """REGIME_GATE.fallback = "best" 时: 全组合未过线 → 返回 best-alpha 组合.

    2026-08-05 用户否决 "best" 兜底 ("如果没有优势就不要入选") — fallback="none"
    时恒返回 None → 主视界无优势即空仓观望, 不输出清单."""
    if REGIME_GATE.get("fallback") != "best":
        return None
    best, best_alpha = None, -1e9
    for (b, s), g in gate.items():
        if not g["per"]:
            continue
        a = max(p["alpha"] for p in g["per"].values())
        if a > best_alpha:
            best, best_alpha = (b, s), a
    return best


def fmt_regime(gate: dict) -> list[str]:
    """制度门判定摘要 (SUMMARY 顶部): 每组合 top vs 池基线 alpha + 保留/剔除."""
    cfg = REGIME_GATE
    lines = [
        "",
        f"[制度自适应门] 板块×系统保留判定 — 基于最新 OOS 市场数据 "
        f"(固定视界净收益 label_pm, top{int(cfg['top_quantile'] * 100)}% vs 池基线, "
        f"门槛 +{cfg['margin'] * 100:.0f}pp, 判定只看主视界 "
        f"T+{cfg['primary_horizon'][:-1]}):",
    ]
    active = []
    for b, s in (
        ("main", "sniper"),
        ("main", "fusion"),
        ("dual", "sniper"),
        ("dual", "fusion"),
    ):
        g = gate.get((b, s))
        if not g or not g["per"]:
            continue
        seg = [
            f"T+{h[:-1]} top {p['top_ev']:+.2%}(wr {p['top_wr']:.0%}) "
            f"vs 基 {p['base_ev']:+.2%} α{p['alpha']:+.2%}"
            for h, p in g["per"].items()
        ]
        if g["active"]:
            active.append(f"{b}/{s}")
        lines.append(
            f"  {b}/{s:<9}: {'; '.join(seg)}  → {'保留' if g['active'] else '剔除'}"
        )
    if not active:
        note = "无"
    else:
        note = ", ".join(active)
    if not REGIME_GATE.get("enable", True):
        lines.append(
            "  制度门已关闭 (enable=False) → 今日全部板块×系统照常出单, 上述 α 仅作参考"
        )
    else:
        lines.append(
            f"  今日过门组合: {note} — 清单全量输出全部候选, "
            f"未过门组合个股已标注 过门=未过"
        )
        fail = [
            (b, s) for (b, s), g in gate.items() if g and g["per"] and not g["active"]
        ]
        if fail:
            btext = "、".join(sorted({BOARD_LABEL.get(b, b) for b, _ in fail}))
            lines.append(
                f"  ⚠ 今日制度门 {cfg['primary_horizon']} 未过: "
                f"{', '.join(f'{b}/{s}' for b, s in fail)} "
                f"→ 不建议今日买入 {btext} 板块股票"
            )
    return lines


def _add_mfe(df: pd.DataFrame) -> pd.DataFrame:
    """按生产口径 (backtest.add_mfe_labels) 补算四视界 MFE 净标签:
    窗口内最高价 / 买入价 - 1 - cost; 尾段缺未来价 → NaN."""
    horizons = (2, 3, 5, 10)
    g = df.groupby("symbol", sort=False)
    exec_px = g["close_hfq"].shift(-1)
    max_off = max(horizons) + 1
    shifts = pd.concat(
        [g["high_hfq"].shift(-off) for off in range(2, max_off + 1)],
        axis=1,
        keys=range(2, max_off + 1),
    )
    slip = df["adv20"].map(slippage_tier)
    cost_total = COST + 2 * slip
    for k in horizons:
        peak = shifts.loc[:, 2 : k + 1].max(axis=1, skipna=False)
        df[f"mfe_{k}d"] = peak / exec_px - 1 - cost_total
    return df


_PANEL_CACHE: dict = {
    "ready": False,
    "data": {},
}  # _panel_per_stock 结果缓存 (被候选池+校准各调一次)


def _panel_per_stock() -> dict[tuple[str, str], pd.DataFrame]:
    """3y 面板 → 每股最近 ~PER_STOCK_WINDOW 交易日日频 (score, mfe) 序列.

    [2026-08-06 用户: 价格预测必须每股独立, 只用该股自己最近 6 个月历史]
    短名单 OOS 文件只含"每日被选股" → 每股仅零星几行, 无法做每股回归;
    面板 (每 stock × 每交易日) 才有逐股日频. 在此按生产口径重算:
      - score = 与生产一致的 6 特征截面分位 (pv_corr_5 面板同样缺列, 自动跳过)
      - mfe   = _add_mfe 口径 (窗口最高价 / 买入价 - 1 - cost, 净)
    返回 {(board, key): DataFrame[symbol, date, score, mfe_3d..mfe_10d]}.
    """
    out: dict[tuple[str, str], pd.DataFrame] = {}
    need = [
        "symbol",
        "date",
        "close_hfq",
        "high_hfq",
        "adv20",
        "label_pm_3d_net",
        "label_pm_5d_net",
        "label_pm_10d_net",
    ] + [c for c in set(SNIPER.pool) | set(FUSION.pool) if c != "pv_corr_5"]
    if _PANEL_CACHE.get("ready"):
        return _PANEL_CACHE["data"]
    for board in ("main", "dual"):
        fp = DATA_DIR / f"_diag_stage_{board}_3y.parquet"
        dates = pd.to_datetime(
            pq.read_table(str(fp), columns=["date"]).to_pandas()["date"]
        )
        uniq = np.unique(dates.values)
        if len(uniq) < PER_STOCK_WINDOW + 12:
            continue
        cutoff = uniq[-(PER_STOCK_WINDOW + 12)]
        t = pq.read_table(
            str(fp), columns=need, filters=[("date", ">=", cutoff)]
        ).to_pandas()
        if t.empty:
            continue
        t["symbol"] = t["symbol"].astype(str)
        t = t.sort_values(["symbol", "date"]).reset_index(drop=True)
        t = _add_mfe(t)
        fr_sn = t[["symbol", "date"]].copy()
        fr_sn["score"] = pool_score(t, SNIPER.pool)
        fr_fu = t[["symbol", "date"]].copy()
        fr_fu["score"] = pool_score(t, FUSION.pool)
        for h in HORIZONS:
            fr_sn[f"mfe_{h}"] = t[f"mfe_{h}"]
            fr_fu[f"mfe_{h}"] = t[f"mfe_{h}"]
        # 共享 calibrate_mag10d 的目标列 (mag_10d 排名键, 2026-08-07 定案) + 平均预测目标
        # (pred_ret_{3,5,10}d = close-to-close 净收益 label_pm_{h}_net, 非 MFE)
        for _h in ("3d", "5d", "10d"):
            fr_sn[f"label_pm_{_h}_net"] = t[f"label_pm_{_h}_net"]
            fr_fu[f"label_pm_{_h}_net"] = t[f"label_pm_{_h}_net"]
        out[(board, "sniper")] = fr_sn
        out[(board, "fusion")] = fr_fu
        # "both" = 合并短名单口径 (build_merged_shortlist: score=max 两系统分位分) →
        # 共现行也启用逐股校准, 不再回退横截面
        fr_bt = fr_sn.copy()
        fr_bt["score"] = np.maximum(fr_sn["score"], fr_fu["score"])
        for _h in ("3d", "5d", "10d"):
            fr_bt[f"label_pm_{_h}_net"] = t[f"label_pm_{_h}_net"]
        out[(board, "both")] = fr_bt
    _PANEL_CACHE["ready"], _PANEL_CACHE["data"] = True, out
    return out


def expand_candidates(full: pd.DataFrame) -> pd.DataFrame:
    """FULL RUN 短名单加宽: 全板块候选 → 按 pred_mag_10d 取每板块 TOP-10.

    2026-08-07 定案: 250d OOS 证明"每股 pred_mag(幅度) 排名"胜过"特征分排名"
    (MFE 全视界 +3.9~10.1pp, 上涨率打平). 原短名单仅 ~15 只 (特征 TOP-5∪TOP-10),
    pred_mag 在其内重排近乎无效 → 候选池加宽到**全板块** (latest 截面全部有分股,
    不再 TOP-30 预筛), 再由 rank_and_truncate 按 pred_mag_10d 取 TOP-10.
    systems/co_occur = 旧 top-N 标签 (狙击 top5 / 融合 top10, 按特征分) 纯元数据标注
    (与并行 build_merged_shortlist 同约定); 两系统 top-N 之外 → systems="", co_occur=False.
    沿用 _panel_per_stock 与生产一致的截面分 (不重算特征); 原"不重选股"约定被本定案取代.
    """
    panel = _panel_per_stock()
    rows = []
    for board in ("main", "dual"):
        sn = panel.get((board, "sniper"))
        fu = panel.get((board, "fusion"))
        if sn is None or fu is None or sn.empty or fu.empty:
            continue
        last = max(sn["date"].max(), fu["date"].max())

        def _cs(pf: pd.DataFrame, last=last) -> pd.DataFrame:
            return pf[pf["date"] == last][["symbol", "score"]].dropna()

        sn_cs, fu_cs = _cs(sn), _cs(fu)
        sscore = sn_cs.set_index("symbol")["score"]
        fscore = fu_cs.set_index("symbol")["score"]
        # 旧 top-N 标签 (元数据): 狙击 top5 / 融合 top10, 按各自特征分
        sset = set(
            sn_cs.sort_values("score", ascending=False).head(SNIPER.top_n)["symbol"]
        )
        fset = set(
            fu_cs.sort_values("score", ascending=False).head(FUSION.top_n)["symbol"]
        )
        # 全板块: latest 截面全部有分股 (不再 TOP-30 预筛)
        for sym in sscore.index.union(fscore.index):
            in_s, in_f = sym in sset, sym in fset
            rows.append(
                {
                    "date": str(pd.Timestamp(last).date()),
                    "board": board,
                    "symbol": sym,
                    "systems": (
                        "fusion+sniper"
                        if in_s and in_f
                        else ("sniper" if in_s else ("fusion" if in_f else ""))
                    ),
                    "score": float(max(sscore.get(sym, -1.0), fscore.get(sym, -1.0))),
                    "co_occur": bool(in_s and in_f),
                    "rk": 0,
                    "cut": "T-5",  # 占位: rank_and_truncate 统一覆盖
                }
            )
    cands = pd.DataFrame(rows)
    base = ["date", "board", "symbol", "systems", "score", "co_occur", "rk", "cut"]
    keep = full[base].copy()
    keep["date"] = keep["date"].astype(str).str.slice(0, 10)
    cands["date"] = cands["date"].astype(str).str.slice(0, 10)
    merged = pd.concat([cands, keep], ignore_index=True)
    merged = merged.drop_duplicates(subset=["board", "symbol"], keep="first")
    print(
        f"[cand] 候选池: 原短名单 {len(full):,} → 全板块 {len(merged):,} "
        f"(latest 截面有分股 {len(cands):,}; 旧 top-N 仅作 systems 标注)",
        flush=True,
    )
    return merged


class _PooledReg:
    """单变量线性映射 (收缩后斜率 + 质心截距), 提供与 LinearRegression 一致的 predict.

    slope = λ·slope_per + (1-λ)·slope_cross (λ=n/(n+SHRINK_KAPPA)),
    intercept = ȳ − slope·x̄ → 拟合线过该股 (score, mfe) 质心.
    """

    __slots__ = ("coef_", "intercept_")

    def __init__(self, slope: float, intercept: float):
        self.coef_ = np.array([slope])
        self.intercept_ = intercept

    def predict(self, X) -> np.ndarray:
        return X @ self.coef_ + self.intercept_


def _fit_calibrators(records: dict) -> dict[tuple, dict]:
    """[2026-08-06] 价格预测=每股独立, 只用该股自己最近 6 个月历史 (用户定案):
    - mag: 每股自己最近 ~PER_STOCK_WINDOW 交易日拟合 LinearRegression (score→mfe),
           数据来自 3y 面板逐股日频 (_panel_per_stock), 非 OOS 稀疏清单;
           该股样本 ≥PER_STOCK_MIN_N 才启用; 不足回退板块×系统横截面回归
    - prob: 横截面 Platt (score → P(mfe_h ≥ ABS_TARGET[h])), 每股分数→唯一概率
    返回 {(board, key, h): {"prob": lr, "mag_cross": reg, "per": {symbol: reg}}}."""
    cals = {}
    panel = _panel_per_stock()
    for (board, key), rec in records.items():
        for h in HORIZONS:
            col = f"mfe_{h}"
            sub = rec[["score", col]].dropna()
            if len(sub) < 20:
                continue
            prob = LogisticRegression()
            prob.fit(
                sub[["score"]].to_numpy(),
                (sub[col] >= ABS_TARGET[h]).astype(int).to_numpy(),
            )
            cross = LinearRegression()
            cross.fit(sub[["score"]].to_numpy(), sub[col].to_numpy())
            per: dict[str, _PooledReg] = {}
            pf = panel.get((board, key))
            if pf is not None:
                for sym, g in pf.groupby("symbol"):
                    gg = (
                        g.dropna(subset=["score", col])
                        .sort_values("date")
                        .tail(PER_STOCK_WINDOW)
                    )
                    if len(gg) >= PER_STOCK_MIN_N:
                        x = gg[["score"]].to_numpy()
                        y = gg[col].to_numpy()
                        raw = LinearRegression().fit(x, y)
                        lam = len(gg) / (len(gg) + SHRINK_KAPPA)
                        slope = lam * float(raw.coef_[0]) + (1 - lam) * float(
                            cross.coef_[0]
                        )
                        per[sym] = _PooledReg(
                            slope, float(y.mean()) - slope * float(x.mean())
                        )
            cals[(board, key, h)] = {"prob": prob, "mag_cross": cross, "per": per}
    return cals


def calibrate(
    records: dict, board: str, key: str, symbol: str, score: float, cals: dict
) -> dict[str, tuple]:
    """最新 score → 逐视界 (预期涨幅, 自然达到概率).

    每股独立 (2026-08-06 用户: 价格预测必须每股独立, 只用该股自己最近 6 个月历史):
    - 预期涨幅 = 该股 score 经该股自己最近 ~6 个月 (score→mfe) 线性回归的独立值,
      该股样本不足时回退板块×系统横截面回归
    - 概率     = P(该股该视界达到固定绝对涨幅 ABS_TARGET) via 横截面 Platt (每股分数→唯一概率)
    返回 {h: (exp_mfe, hit_prob)}; 无校准数据 → (NaN, NaN).
    """
    out = {h: (float("nan"), float("nan")) for h in HORIZONS}
    rec = records.get((board, key))
    if rec is None or len(rec) == 0:
        return out
    for h in HORIZONS:
        c = cals.get((board, key, h))
        if c is None:
            continue
        mag = c["per"].get(symbol, c["mag_cross"])
        exp = float(mag.predict(np.array([[score]]))[0])
        prob = float(c["prob"].predict_proba(np.array([[score]]))[0, 1])
        out[h] = (exp, prob)
    return out


def _c2c_latest(panel: dict, board: str, h: str, last) -> pd.Series:
    """共享 calibrate_mag10d (score→label_pm_{h}_net) 的决策日每股 close-to-close 平均预期.

    h ∈ ("3d","5d","10d"). 排名键 pred_mag_10d 与并行 build_merged_shortlist 同源
    (2026-08-07 d10 c2c 定案); 平均预测 pred_ret_{h} 复用同机制 (2026-08-09 用户:
    看板显示平均预测而非 MFE 最大). 用合并 "both" 面板 (score=max 两系统池分),
    calibrate_mag10d 内部已做无前瞻 (已实现边界 buy_lag+label_horizon 交易日, 按视界).
    last 必须等于该板块面板最新交易日 (= 决策日). 返回 symbol→mag Series.
    """
    fr = panel.get((board, "both"))
    if fr is None or fr.empty or f"label_pm_{h}_net" not in fr.columns:
        return pd.Series(dtype=float)
    # _panel_per_stock 的帧不含 board 列 → 补上 (该帧本就是单板块, 赋常量 board 正确)
    work = fr[["symbol", "date", "score", f"label_pm_{h}_net"]].copy()
    work["board"] = board
    m = calibrate_mag10d(
        work, target_col=f"label_pm_{h}_net", label_horizon=int(h[:-1])
    )
    if m.empty:
        return pd.Series(dtype=float)
    row = m[m["date"] == last]
    if row.empty:
        return pd.Series(dtype=float)
    return row.set_index("symbol")["mag"]


def _mag10d_latest(panel: dict, board: str, last) -> pd.Series:
    """排名键 pred_mag_10d: 委托 _c2c_latest (score→label_pm_10d_net, T+10 c2c)."""
    return _c2c_latest(panel, board, "10d", last)


def _anchor_frame(board: str, window: int = ANCHOR_WINDOW) -> pd.DataFrame:
    """报告锚专用加宽面板 (每板块 ~window+12 交易日 score + close-to-close 净收益).

    预测面板 _panel_per_stock 只留 ~PER_STOCK_WINDOW+12=142 交易日 (每股 6 个月语义),
    装不下 250d 已实现锚; 此处独立读取同款 _diag_stage parquet, 同公式算 score
    (max 两系统池分), 供 _trailing_realized 求 top-ANCHOR_TOP 已实现均值. score 是
    跨日横截面分位, 与预测面板重叠日完全一致 → 锚定的"近窗 top-N"就是模型当日入选集.
    """
    fp = DATA_DIR / f"_diag_stage_{board}_3y.parquet"
    dates = pd.to_datetime(pq.read_table(str(fp), columns=["date"]).to_pandas()["date"])
    uniq = np.unique(dates.values)
    if len(uniq) < window + 12:
        return pd.DataFrame()
    cutoff = uniq[-(window + 12)]
    pool_cols = [c for c in set(SNIPER.pool) | set(FUSION.pool) if c != "pv_corr_5"]
    cols = ["symbol", "date"] + pool_cols + [f"label_pm_{h}_net" for h in HORIZONS]
    t = pq.read_table(
        str(fp), columns=cols, filters=[("date", ">=", cutoff)]
    ).to_pandas()
    if t.empty:
        return pd.DataFrame()
    t["symbol"] = t["symbol"].astype(str)
    sn = pool_score(t, SNIPER.pool)
    fu = pool_score(t, FUSION.pool)
    t["score"] = np.maximum(sn.values, fu.values)
    keep = ["symbol", "date", "score"] + [f"label_pm_{h}_net" for h in HORIZONS]
    return t[keep].dropna(subset=["score"]).reset_index(drop=True)


def _trailing_realized(
    frame: pd.DataFrame, h: str, top: int = ANCHOR_TOP, window: int = ANCHOR_WINDOW
) -> float:
    """模型近 window 个已实现决策日的 top-top 已实现净收益均值 (报告幅度锚).

    面板 (board,"both") 每股日频 score + label_pm_{h}_net; 每决策日按 score 降序取
    top 只 (score 与 pred_mag 单调 → 即模型当日入选集), 求 label 均值; 只统计已实现日
    (label 非 NaN). 返回该板块最近 window 个已实现决策日的 top-top 均值.
    """
    lab = f"label_pm_{h}_net"
    df = frame.dropna(subset=[lab])
    if df.empty:
        return float("nan")
    days = sorted(df["date"].unique())[-window:]
    sub = df[df["date"].isin(days)]
    top_rows = (
        sub.sort_values(["date", "score"], ascending=[True, False])
        .groupby("date", sort=False)
        .head(top)
    )
    return float(top_rows[lab].mean())


def _anchor_reported(res: pd.DataFrame, panel: dict | None = None) -> pd.DataFrame:
    """报告幅度锚定 (2026-08-14 用户: 并行预测高于实际, 主板+双创都要).

    根因: pred_ret_{h} = 短窗(cal_n=21)横截面 OLS 对今日热 score 外推 → 顺周期幅度高估
    (诊断 _diag_mag10d_walkforward/_diag_cal_window_compare: 250d 双创 top10 预测 +4.68%
    实得 +5.56% 反而低估, 但热行情单日预测飙到 +12~13% 是 2-3x 高估; 拉长校准窗到 125d
    更差: 近期 120d bias +2.84pp/实得 +0.14% vs cal_n=21 +0.23pp/+3.46%). 裁决: 排名键
    cal_n=21 保留 (排序有真本事), 报告幅度改用模型近 ANCHOR_WINDOW 决策日 top-ANCHOR_TOP
    已实现均值做水平锚 — 每板块每视界整体平移, 保持日内横截面差异与排序不变, 只把整体
    水平拉回诚实位置 (2026-08-14: 250d 双创 top-5 实得 ~+6.3%, 主板 ~+3.5%; 用户"3% 低估"
    成立, 120d 短窗恰踩双创弱行情). pred_mag_10d 同步 (docx/合并表展示键).

    panel 为测试注入面 (dict[(board,"both")]→DataFrame); 生产不传, 走 _anchor_frame
    独立加宽读取 (预测面板仅 142 交易日, 装不下 250d 锚).
    """
    out = res.copy()
    for board in ("main", "dual"):
        if panel is not None:
            fr = panel.get((board, "both"))
        else:
            fr = _anchor_frame(board)
        if fr is None or fr.empty:
            continue
        today = out[out["board"] == board]
        if today.empty:
            continue
        t5 = today[today["cut"] == "T-5"]
        if t5.empty:
            t5 = today
        for h in HORIZONS:
            col = f"pred_ret_{h}"
            if col not in out.columns or f"label_pm_{h}_net" not in fr.columns:
                continue
            t_real = _trailing_realized(fr, h, top=ANCHOR_TOP, window=ANCHOR_WINDOW)
            if not np.isfinite(t_real):
                continue
            t_pred = float(t5[col].mean())
            shift = t_pred - t_real
            mask = (out["board"] == board) & out[col].notna()
            out.loc[mask, col] = out.loc[mask, col] - shift
            print(
                f"[anchor] {board} T+{h[:-1]} 报告均值 {t_pred:+.1%} → "
                f"实得锚 {t_real:+.1%} (平移 {shift:+.1%})",
                flush=True,
            )
    # 排名键与 docx 头条同步: pred_mag_10d == pred_ret_10d (同源 score→label_pm_10d_net)
    if "pred_mag_10d" in out.columns and "pred_ret_10d" in out.columns:
        m = out["pred_ret_10d"].notna()
        out.loc[m, "pred_mag_10d"] = out.loc[m, "pred_ret_10d"]
    return out


def add_oos_pred(res: pd.DataFrame, records: dict) -> pd.DataFrame:
    """每只短名单股: 用最新 score 经 OOS 校准给 逐视界 前瞻 预期涨幅(MFE)+逐股自然概率.

    2026-08-07: 排名键 pred_mag_10d 用共享 calibrate_mag10d (score→label_pm_10d_net,
    T+10 close-to-close 校准幅度) 覆盖 — 与并行 build_merged_shortlist 同源; 其余
    pred_mag_{2,3,5}d 与全部 pred_prob 保持每股 MFE/自然概率口径不变 (prob 校准与
    select_confident T+2/T+3 门不动).
    2026-08-09: 另输出 pred_ret_{3,5,10}d = 每股 close-to-close 平均预期 (score→
    label_pm_{h}_net, 非 MFE 最大), 供看板"预期"列展示 (用户定案).
    """
    cals = _fit_calibrators(records)
    out = res.copy()
    for h in HORIZONS:
        out[f"pred_mag_{h}"] = float("nan")
        out[f"pred_prob_{h}"] = float("nan")
    for idx, r in out.iterrows():
        key = r["systems"] if r["systems"] in ("sniper", "fusion") else "both"
        cal = calibrate(
            records, r["board"], key, str(r["symbol"]), float(r["score"]), cals
        )
        for h in HORIZONS:
            mag, prob = cal[h]
            out.at[idx, f"pred_mag_{h}"] = mag
            out.at[idx, f"pred_prob_{h}"] = prob
    # 排名键覆盖: pred_mag_10d ← 共享 calibrate_mag10d (无前瞻, 全板块日截面)
    panel = _panel_per_stock()
    for board in ("main", "dual"):
        fr = panel.get((board, "both"))
        if fr is None or fr.empty:
            continue
        mag10 = _mag10d_latest(panel, board, fr["date"].max())
        if mag10.empty:
            continue
        mask = out["board"] == board
        out.loc[mask, "pred_mag_10d"] = out.loc[mask, "symbol"].map(mag10)
    # 平均预期 (close-to-close, 非 MFE): 每视界 每股 score→label_pm_{h}_net 校准
    # (2026-08-09 用户: 看板显示平均预测而非 MFE 最大). 与排名键同源; pred_mag 保留不动.
    for h in HORIZONS:
        out[f"pred_ret_{h}"] = float("nan")
    for board in ("main", "dual"):
        fr = panel.get((board, "both"))
        if fr is None or fr.empty:
            continue
        mask = out["board"] == board
        last = fr["date"].max()
        for h in HORIZONS:
            mag = _c2c_latest(panel, board, h, last)
            if mag.empty:
                continue
            out.loc[mask, f"pred_ret_{h}"] = out.loc[mask, "symbol"].map(mag)
    # est_wr 已移除: 系统级常量 (同系统每股同值), 且是旧 P(MFE>0) 虚高口径 (用户 2026-08-05 否决)
    first = ["date", "board", "cut", "rk", "symbol", "systems", "co_occur", "score"]
    ph = [f"{k}_{h}" for h in HORIZONS for k in ("pred_mag", "pred_prob")]
    pr = [f"pred_ret_{h}" for h in HORIZONS]
    return out[first + ph + pr].reset_index(drop=True)


def _module_suffix(module: str) -> str:
    return f"__{module}" if module != "na" else ""


def _persist_raw_preds(res: pd.DataFrame, sel_date: pd.Timestamp, module: str) -> None:
    """WORM: 每股当日 raw 预测 (平滑前) → parallel_preds_raw_<date>__<module>.csv.

    EMA 平滑的历史底稿; 旧文件不覆盖. 同 symbol 多 cut 行去重 (预测值相同, keep 最后一行).
    """
    if not SMOOTH_ENABLED:
        return
    os.makedirs(str(STOCK_LIST_DIR), exist_ok=True)
    cols = ["date", "board", "symbol", "systems", "score"] + [
        f"{k}_{h}" for h in HORIZONS for k in ("pred_mag", "pred_prob")
    ]
    d = res[cols].copy()
    d["date"] = str(sel_date.date())
    d = d.drop_duplicates(subset=["symbol"], keep="last").reset_index(drop=True)
    stamp = str(sel_date.date()).replace("-", "")
    fp = STOCK_LIST_DIR / f"parallel_preds_raw_{stamp}{_module_suffix(module)}.csv"
    d.to_csv(fp, index=False)
    print(f"[raw] {fp}", flush=True)


def _load_raw_history(sel_date: pd.Timestamp, module: str) -> pd.DataFrame:
    """读 <选股日> 之前的 raw 预测文件 → 长表 (symbol, hist_date, pred_mag/prob_{h})."""
    suffix = _module_suffix(module)
    today = str(sel_date.date()).replace("-", "")
    pred_cols = [f"{k}_{h}" for h in HORIZONS for k in ("pred_mag", "pred_prob")]
    frames = []
    for fp in STOCK_LIST_DIR.glob("parallel_preds_raw_*.csv"):
        m = re.match(r"parallel_preds_raw_(\d{8})(.*)\.csv$", fp.name)
        if not m:
            continue
        if suffix and m.group(2) != suffix:
            continue
        if not suffix and m.group(2):
            continue  # module=na 时不混入带模块标记的历史
        if m.group(1) >= today:
            continue  # 不含今日 (今日 raw 由调用方直接给)
        d = pd.read_csv(fp, dtype={"symbol": str})
        d = d[["symbol"] + pred_cols]
        d["hist_date"] = m.group(1)
        frames.append(d)
    if not frames:
        return pd.DataFrame()
    long = pd.concat(frames, ignore_index=True)
    return long.drop_duplicates(subset=["symbol", "hist_date"], keep="last")


def ema_smooth(res: pd.DataFrame, sel_date: pd.Timestamp, module: str) -> pd.DataFrame:
    """输出级时间平滑: 每股今日 pred_mag/prob = 近 SMOOTH_K 个可用交易日 raw 的衰减加权均值.

    w_k = α·(1-α)^k (k=0=今日, 越旧越轻), 截断后归一化; 缺日跳过 (gap-robust).
    只平滑 pred_mag/pred_prob (不动 score); 无历史 → 保持 raw. 直接抑制相邻交易日预测/概率剧变.
    """
    if not SMOOTH_ENABLED:
        return res
    hist = _load_raw_history(sel_date, module)
    if hist.empty:
        return res
    weights = np.array(
        [SMOOTH_ALPHA * (1 - SMOOTH_ALPHA) ** k for k in range(SMOOTH_K)]
    )
    weights /= weights.sum()
    out = res.copy()
    for sym in out["symbol"].unique():
        h = hist[hist["symbol"] == sym].sort_values("hist_date", ascending=False)
        if h.empty:
            continue
        src = h.head(SMOOTH_K - 1)  # 最多 SMOOTH_K-1 个旧日 (今日占 k=0)
        mask = out["symbol"] == sym
        for kind in ("pred_mag", "pred_prob"):
            for hz in HORIZONS:
                col = f"{kind}_{hz}"
                today_v = out.loc[mask, col]
                if today_v.empty or not np.isfinite(float(today_v.iloc[0])):
                    continue
                pairs = [
                    (w, v)
                    for w, v in zip(
                        weights,
                        [float(today_v.iloc[0])] + [float(x) for x in src[col]],
                        strict=False,
                    )
                    if np.isfinite(v)
                ]
                if not pairs:
                    continue
                ww = np.array([p[0] for p in pairs])
                ww /= ww.sum()
                out.loc[mask, col] = float(np.dot([p[1] for p in pairs], ww))
    return out


def _sel_reason(r: pd.Series) -> str:
    p = "n/a" if pd.isna(r["pred_prob_3d"]) else f"{r['pred_prob_3d']:.0%}"
    m3 = "n/a" if pd.isna(r["pred_ret_3d"]) else f"{r['pred_ret_3d']:+.1%}"
    return f"{r['symbol']}(T+3 {m3}/{p})"


def select_confident(res: pd.DataFrame, prob_min: float = 0.0) -> pd.DataFrame:
    """IRON RULE (用户): 只列预测上涨股. 入选门 = 纯 T+3 门.

    2026-08-09 删 2d 视界: 原 T+2/T+3 联合门 (2026-08-07) 退化为纯 T+3 —
    2d 视界及其 pred_mag_2d 不再存在, 副门 (T+2 强看涨) 随之删除 (config
    SHORTLIST_SCORE.select_gate):
      保留 ⇔ T+3 可兑现净预期涨幅 (pred_ret_3d, close-to-close) > t3_min

    2026-08-10: 门从 pred_mag_3d (MFE 最大浮盈, 虚高不可兑现) 改为 pred_ret_3d
    (label_pm_3_net close-to-close 净预期, 成本已扣) — 与 legacy 收益闸一致.
    2026-08-14: t3_min 分板块 dict (main=0 / dual=0.5%, _diag_t3min_sweep 定案),
    也可为全局 float (向后兼容).

    概率口径 (用户 2026-08-06): 概率=逐股自然概率 (P(该股达到固定绝对目标)), 每股唯一
    真值. 原 ">60%" 门槛 (基于 P(MFE>0)≈90% 旧口径) 不可达, 故默认不设概率门槛.
    保留 prob_min 参数以便后续收紧. 主视界 T+3 (2026-08-05 用户: 短持 3 天)."""
    g3 = res["pred_ret_3d"]
    sg = SHORTLIST_SCORE.get("select_gate", {})
    t3_min = sg.get("t3_min", 0.0)
    # 2026-08-14: t3_min 支持分板块 dict (main=0 / dual=0.5%), 或全局 float (向后兼容)
    if isinstance(t3_min, dict):
        thr = res["board"].map(t3_min).fillna(0.0)
        keep = g3 > thr
    else:
        thr = t3_min
        keep = g3 > t3_min
    if prob_min > 0:
        keep = keep & (res["pred_prob_3d"] > prob_min)
    dropped = res[~keep]
    if len(dropped):
        prob_txt = "" if prob_min <= 0 else f" 或 达到概率≤{prob_min:.0%}"
        gate_txt = (
            "、".join(f"{b}={v:.1%}" for b, v in t3_min.items())
            if isinstance(t3_min, dict)
            else f"{t3_min:.1%}"
        )
        print(
            f"[select] 剔除 {len(dropped)} 只 (T+3≤{gate_txt}{prob_txt}): "
            f"{', '.join(dropped.apply(_sel_reason, axis=1))}",
            flush=True,
        )
    return res[keep].copy().reset_index(drop=True)


def add_score(df: pd.DataFrame) -> pd.DataFrame:
    """按 SHORTLIST_SCORE 权重合成综合分 score_w (用户: 预测先行→按权重打分→再排名).

    score_w = Σ_h horizon_w[h] × (gain_w×norm_gain_h + prob_w×norm_prob_h)
    涨幅与达到概率都按当日入选股 min-max 归一化到 0~1 (同量纲, 真 50:50).
    2026-08-05 用户定案: 概率若用原始值 (0.32~0.40 近似常数), 排名被涨幅主导 (相关 0.949);
    归一化后概率真正参与排名. 排名=score_w 降序, 绝不先排名后预测.
    """
    out = df.copy()
    for h in HORIZONS:
        g = out[f"pred_mag_{h}"]
        glo, ghi = g.min(), g.max()
        out[f"norm_g_{h}"] = ((g - glo) / (ghi - glo)).fillna(0.0) if ghi > glo else 0.0
        p = out[f"pred_prob_{h}"].fillna(0.0)
        plo, phi = p.min(), p.max()
        out[f"norm_p_{h}"] = ((p - plo) / (phi - plo)) if phi > plo else 0.0
    hw, gw, pw = (
        SHORTLIST_SCORE["horizon_w"],
        SHORTLIST_SCORE["gain_w"],
        SHORTLIST_SCORE["prob_w"],
    )
    out["score_w"] = sum(
        hw[h] * (gw * out[f"norm_g_{h}"] + pw * out[f"norm_p_{h}"]) for h in HORIZONS
    )
    return out


def rank_and_truncate(res: pd.DataFrame) -> pd.DataFrame:
    """2026-08-14 定案: 入选收紧至每板块 TOP-5.

    250d OOS 幅度前沿 (诊断 _diag_mag_frontier): top5 幅度/命中率/大涨率全优于 top10
    (双创实得 +6.34% vs +5.76%, 主板 +3.49% vs +3.22%), 只牺牲每日出股数 10→5.
    按 CAND_RANK_KEY (pred_mag_10d) 每板块降序取前 5, cut 统一标 T-5.
    """
    key = CAND_RANK_KEY
    if res.empty:
        return res
    out = []
    for board in ("main", "dual"):
        b = res[res["board"] == board].sort_values(
            key, ascending=False, na_position="last"
        )
        if b.empty:
            continue
        out.append(b.head(5).assign(cut="T-5"))
    return pd.concat(out, ignore_index=True).reset_index(drop=True)


def build_merged(res: pd.DataFrame) -> pd.DataFrame:
    """合并短名单: 每板块 TOP-5 (入选档); main+dual 全局排名.

    IRON RULE (用户): 先预测(涨幅+达到概率) → 再排名.
    排名=pred_mag_3d 降序 (2026-08-07 定案: 纯幅度排名 > score_w 混合), 平局按 score_w→
    达到概率; 负涨幅绝不在正涨幅之前 (已由 select 门保证). 共现仅作平局裁决参考.
    """
    df = res[res["cut"] == "T-5"].copy()
    t5 = set(df["symbol"])
    df["in_t5"] = df["symbol"].isin(t5)
    df = add_score(df)
    keys = [CAND_RANK_KEY, "score_w", "pred_prob_3d"]
    df = df.sort_values(keys, ascending=False, na_position="last").reset_index(
        drop=True
    )
    df.insert(0, "rank", range(1, len(df) + 1))
    return df


def fmt_merged(m: pd.DataFrame) -> list[str]:
    """合并排名表的终端文本 (单张决策清单, 涨幅→概率降序, 前瞻逐视界)."""
    hdr = (
        f"{'#':>3} {'代码':<8} {'板块':<5} {'模块':<16} {'score':>7} "
        + "  ".join(
            f"{'T+' + h[:-1] + ' 预期涨幅(达到概率)':>20}" for h in HORIZON_ORDER
        )
        + f" {'T5':>3}"
    )
    if len(m) == 0:
        return [
            "── 合并短名单 (主板+双创合一, 共 0 只) ──",
            "  今日主视界无优势组合, 空仓观望 (见下方制度门判定)",
        ]
    lines = [
        f"── 合并短名单 (主板+双创合一, 共 {len(m)} 只; 仅正预期涨幅; "
        f"预期=该股最新score经OOS每股独立线性校准的今后涨幅(MFE), 每股唯一; "
        f"达到概率=逐股自然概率 P(该股达到固定绝对目标), 每股唯一真值, 非桶级共享; "
        f"过门=该板块×系统今日过10d制度门, "
        f"未过门个股已标注 过门=未过, 不建议买入) ──",
        hdr,
    ]
    for _, r in m.iterrows():
        mod = "★共现(双系统)" if bool(r["co_occur"]) else str(r["systems"])
        cols = []
        for h in HORIZON_ORDER:
            ex = (
                "n/a"
                if pd.isna(r[f"pred_mag_{h}"])
                else f"{r[f'pred_mag_{h}']:+.1%}({r[f'pred_prob_{h}']:.0%})"
            )
            cols.append(f"{ex:>18}")
        t5 = "Y" if bool(r["in_t5"]) else ""
        lines.append(
            f"{r['rank']:>3} {r['symbol']:<8} {BOARD_LABEL.get(r['board'], r['board']):<5} "
            f"{mod:<16} {r['score']:.4f} " + "  ".join(cols) + f" {t5:>3}"
        )
    return lines


def build_summary(res: pd.DataFrame, stats: dict, sel_date: pd.Timestamp) -> list[str]:
    """SUMMARY: 系统级逐视界胜率/期望 → 两系统共识 → 建议顺序."""
    lines = [
        f"── SUMMARY ── 选股日(数据) {sel_date:%Y-%m-%d}",
        "系统级 OOS 逐视界 期望涨幅(MFE)/达到概率 (同系统个股共享, 非个股值; 真实口径):",
    ]
    for board in ("main", "dual"):
        label = BOARD_LABEL.get(board, board)
        lines.append(f"  [{label}]")
        for name in ("sniper", "fusion"):
            per = stats.get((board, name), {})
            seg = "  ".join(
                f"T+{h[:-1]} {per[h]['mag']:+.1%}({per[h]['hit']:.0%})"
                for h in HORIZONS
                if h in per
            )
            lines.append(f"    {name}: {seg or 'n/a'}")
    for board in ("main", "dual"):
        b = res[res["board"] == board]
        if b.empty:
            continue
        label = BOARD_LABEL.get(board, board)
        co = b[b["co_occur"]].sort_values(
            ["cut", CAND_RANK_KEY], ascending=[True, False]
        )
        lines.append(
            f"\n[{label}] 两系统共识(共现)股: "
            f"{', '.join(str(s) for s in co['symbol'].unique()) if not co.empty else '无'}"
        )
        lines.append(f"  建议顺序 ({CAND_RANK_KEY} 降序: 主视界每股预期幅度):")
        for cut in ("T-5",):
            g = b[b["cut"] == cut].sort_values(CAND_RANK_KEY, ascending=False)
            if g.empty:
                continue
            picks = " > ".join(
                f"{r['symbol']}#{i + 1}[{r['systems']}]{'★' if bool(r['co_occur']) else ''}"
                for i, (_, r) in enumerate(g.iterrows())
            )
            lines.append(f"    {cut}: {picks}")
    lines.append(
        "\n每只个股逐视界预期 = 该股最新 score 经 OOS 每股独立线性校准的 今后涨幅(MFE) "
        "(每股唯一, 非桶级共享; 前瞻, 非历史回看)."
    )
    lines.append(
        "达到概率 = 逐股自然概率 P(该股达到固定绝对目标) — 每股唯一真值, 非桶级共享 "
        "(旧 P(MFE>0) 85-99% 已废, 2026-08-05 用户定案)."
    )
    lines.append(
        "建议: 优先 score_w 高者 (预期涨幅×概率加权); 系统级期望/胜率见 SUMMARY 段."
    )
    return lines


def write_docx(
    res: pd.DataFrame,
    summary: list[str],
    sel_date: pd.Timestamp,
    path: Path,
    merged: pd.DataFrame | None = None,
    module: str = "",
    rejected: dict[str, str] | None = None,
) -> None:
    if Document is None:
        return
    doc = Document()
    title = f"STOCK LIST {sel_date:%Y%m%d}"
    if module:
        title += f" · module {module}"
    doc.add_heading(title, level=0)
    if merged is not None and not merged.empty:
        doc.add_heading("合并短名单 · 集成决策清单", level=1)
        doc.add_paragraph(
            "rank=score_w 降序 (预期涨幅×达到概率加权, T+3 主视界短持) · "
            "systems=命中模块(共现=双系统) · T5=该股是否入选所在板块T-5 · 仅保留 T+3 预期涨幅>0 的股 · "
            "每视界(T+2/3/5/10) 预期涨幅(MFE)=最新score经OOS每股独立线性校准的今后表现, 每股唯一; 达到概率=逐股自然概率 P(该股达到固定绝对目标), 每股唯一真值"
        )
        mcols = [
            "rank",
            "symbol",
            "board",
            "module",
            "score",
            "score_w",
            "in_t5",
            "过门",
        ] + [f"{k}_{h}" for h in HORIZON_ORDER for k in ("pred_mag", "pred_prob")]
        t = doc.add_table(rows=1, cols=len(mcols))
        for j, c in enumerate(mcols):
            t.rows[0].cells[j].text = c
        for _, r in merged.iterrows():
            cells = t.add_row().cells
            cells[0].text = str(r["rank"])
            cells[1].text = str(r["symbol"])
            cells[2].text = BOARD_LABEL.get(r["board"], r["board"])
            cells[3].text = str(r["systems"])
            cells[4].text = f"{float(r['score']):.4f}"
            cells[5].text = f"{float(r['score_w']):.4f}"
            cells[6].text = "Y" if bool(r["in_t5"]) else ""
            cells[7].text = str(r["过门"])
            j = 8
            for h in HORIZON_ORDER:
                cells[j].text = (
                    "n/a"
                    if pd.isna(r[f"pred_mag_{h}"])
                    else f"{float(r[f'pred_mag_{h}']):+.1%}"
                )
                cells[j + 1].text = (
                    "n/a"
                    if pd.isna(r[f"pred_prob_{h}"])
                    else f"{float(r[f'pred_prob_{h}']):.0%}"
                )
                j += 2
    for line in summary:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9)
    # 未接受板块 (被退回) 仍出分板节, 醒目标注原因 — 不静默跳过
    for board, reason in (rejected or {}).items():
        doc.add_heading(f"{board} ({BOARD_LABEL.get(board, board)})", level=1)
        p = doc.add_paragraph()
        run = p.add_run(f"⚠ 未接受 (被退回): {reason} — 当日短名单为空, 未出股")
        run.font.size = Pt(10)
        run.bold = True
    cols = [
        "rank",
        "symbol",
        "module",
        "co_occur",
        "score",
        "score_w",
        "过门",
    ] + [f"{k}_{h}" for h in HORIZONS for k in ("pred_mag", "pred_prob")]
    for board in ("main", "dual"):
        b = res[res["board"] == board]
        if b.empty:
            continue
        doc.add_heading(f"{board} ({BOARD_LABEL.get(board, board)})", level=1)
        for cut in ("T-5",):
            g = b[b["cut"] == cut].sort_values(CAND_RANK_KEY, ascending=False)
            if g.empty:
                continue
            doc.add_paragraph(
                f"{cut} · 按 10d 预期幅度降序(全局质量排名, 非板块内排名) · 仅正预期涨幅股 · "
                f"预期涨幅(MFE)=最新score经OOS每股独立线性校准的今后表现, 每股唯一; 达到概率=逐股自然概率 P(该股达到固定绝对目标)"
            )
            t = doc.add_table(rows=1, cols=len(cols))
            for j, c in enumerate(cols):
                t.rows[0].cells[j].text = c
            for _i, (_, r) in enumerate(g.iterrows()):
                cells = t.add_row().cells
                cells[0].text = str(r["rank"])
                cells[1].text = str(r["symbol"])
                cells[2].text = str(r["systems"])
                cells[3].text = "★" if bool(r["co_occur"]) else ""
                cells[4].text = f"{float(r['score']):.4f}"
                cells[5].text = f"{float(r['score_w']):.4f}"
                cells[6].text = str(r["过门"])
                for j, h in enumerate(HORIZONS):
                    cells[7 + 2 * j].text = (
                        "n/a"
                        if pd.isna(r[f"pred_mag_{h}"])
                        else f"{float(r[f'pred_mag_{h}']):+.1%}"
                    )
                    cells[8 + 2 * j].text = (
                        "n/a"
                        if pd.isna(r[f"pred_prob_{h}"])
                        else f"{float(r[f'pred_prob_{h}']):.0%}"
                    )
    doc.save(str(path))


def write_xlsx(
    res: pd.DataFrame,
    summary: list[str],
    sel_date: pd.Timestamp,
    path: Path,
    merged: pd.DataFrame | None = None,
    module: str = "",
) -> None:
    if Workbook is None:
        return
    wb = Workbook()
    hdr_fill = PatternFill("solid", fgColor="D9E1F2")
    bold = Font(bold=True)
    pct = "0.0%"

    def _sheet(ws, df, cols, pct_cols=()):
        ws.append(cols)
        for c in ws[1]:
            c.font, c.fill = bold, hdr_fill
        for _, r in df.iterrows():
            ws.append([r.get(c, "") for c in cols])
        for j, c in enumerate(cols, 1):
            if c in pct_cols:
                for cell in ws.iter_rows(min_row=2, min_col=j, max_col=j):
                    v = cell[0].value
                    if isinstance(v, (int, float)):
                        cell[0].number_format = pct
            ws.column_dimensions[get_column_letter(j)].width = max(
                10, min(28, max(len(str(c)), 6))
            )
        ws.freeze_panes = "A2"

    ws = wb.active
    ws.title = "SUMMARY"
    title = f"STOCK LIST {sel_date:%Y%m%d}"
    if module:
        title += f" · module {module}"
    ws.append([title])
    ws["A1"].font = Font(bold=True, size=13)
    for line in summary:
        ws.append([line])
    ws.column_dimensions["A"].width = 110

    data_cols = [
        "date",
        "board",
        "cut",
        "rank",
        "symbol",
        "module",
        "co_occur",
        "score",
        "过门",
    ] + [f"{k}_{h}" for h in HORIZONS for k in ("pred_mag", "pred_prob")]
    pct_cols = [f"{k}_{h}" for h in HORIZONS for k in ("pred_mag", "pred_prob")]
    if merged is not None and not merged.empty:
        mcols = [
            "rank",
            "symbol",
            "board",
            "module",
            "score",
            "score_w",
            "in_t5",
            "过门",
        ] + [f"{k}_{h}" for h in HORIZON_ORDER for k in ("pred_mag", "pred_prob")]
        m = merged.copy().rename(columns={"systems": "module"})
        _sheet(
            wb.create_sheet("合并排名"),
            m[mcols],
            mcols,
            pct_cols=[
                f"{k}_{h}" for h in HORIZON_ORDER for k in ("pred_mag", "pred_prob")
            ],
        )
    for cut, title in (("T-5", "短名单 T-5"),):
        r2 = res[res["cut"] == cut].copy().rename(columns={"systems": "module"})
        _sheet(wb.create_sheet(title), r2, data_cols, pct_cols)
    wb.save(str(path))


def main() -> int:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    args = sys.argv[1:]
    trade_date = args[0] if (args and len(args[0]) == 8 and args[0].isdigit()) else None
    watch = [a for a in args if a != trade_date]

    global FULLRUN_DIR
    if trade_date:
        # 自动化集成: 只交付"当日"并行 run_dir; 无当日 run_dir → 大声失败 (并行未跑, 无数据可交付)
        d = _latest_fullrun_dir(prefix=trade_date)
        if d is None:
            print(
                f"[error] 未找到 {trade_date} 当日并行 run_dir, 拒绝交付旧数据",
                flush=True,
            )
            return 1
        FULLRUN_DIR = d

    # 读 FULL RUN 落盘短名单 (权威选股, 不重算); 某板短名单为空 = 该板当日被拒/未接受,
    # 仍出该板清单但标注原因 (不静默跳过, 也不整体失败).
    frames = []
    rejected: dict[str, str] = {}
    for board in ("main", "dual"):
        fp = FULLRUN_DIR / f"shortlist_{board}.csv"
        if not fp.exists():
            reason = _reject_reason(FULLRUN_DIR, board)
            rejected[board] = reason
            print(f"[rejected] {board} 未接受: {reason}", flush=True)
            continue
        df = pd.read_csv(fp, dtype={"symbol": str})
        frames.append(df)
    if not frames:
        print("[error] FULL RUN 短名单为空", flush=True)
        return 1
    full = pd.concat(frames, ignore_index=True)
    sel_date = (
        pd.Timestamp(trade_date) if trade_date else pd.Timestamp(full["date"].max())
    )

    t0 = pd.Timestamp.now()
    from app.pipeline1.model_meta import load_modules, module_id

    module = module_id(load_modules())
    records = load_oos_records()
    gate = regime_gate(records)
    active = [f"{b}/{s}" for (b, s), g in gate.items() if g["active"]]
    print(
        f"[regime] 今日保留组合: {', '.join(active) if active else '无 (主视界无优势, 空仓)'}",
        flush=True,
    )
    cands = expand_candidates(full)
    raw_res = add_oos_pred(cands, records)
    _persist_raw_preds(raw_res, sel_date, module)
    res = select_confident(ema_smooth(raw_res, sel_date, module), prob_min=0.0)
    # 制度门 (2026-08-09 用户: STOCK LIST 显示 ALL CANDIDATES — 不整组剔除; 仍按
    # 10d 门计算, 每行标注 过门=是/未过; 未过门个股不建议买入, 报告里写明)
    if REGIME_GATE.get("enable", True):
        active_mask = res.apply(lambda r: _row_active(r, gate), axis=1)
        res = res.copy()
        res["regime_active"] = active_mask
        res["过门"] = res["regime_active"].map({True: "是", False: "未过"})
        n_fail = int((~active_mask).sum())
        if n_fail:
            print(
                f"[regime] 制度门未过 {n_fail} 行 → 已标注 过门=未过 "
                f"(清单全量输出, 未过门个股不建议买入)",
                flush=True,
            )
    res = add_score(res)
    res = rank_and_truncate(res)
    # 报告幅度锚定 (2026-08-14): 排名键 cal_n=21 保留, 报告 pred_ret_{h}/pred_mag_10d
    # 平移至模型近 ANCHOR_WINDOW 决策日 top-ANCHOR_TOP 已实现均值 — 每板块每视界常数, 排序不变
    res = _anchor_reported(res)
    stats = load_system_stats(records)
    merged = build_merged(res)
    # 全局质量排名 (2026-08-09 用户: 单一清单按预测质量排序, 不按板块分组;
    # rank = 每板块 TOP-5 全池按 CAND_RANK_KEY(pred_mag_10d) 降序 1..10)
    rank_map = merged.set_index("symbol")["rank"]
    res["rank"] = res["symbol"].map(rank_map)
    cut_ord = {"T-5": 0}
    res = (
        res.assign(_co=res["cut"].map(cut_ord))
        .sort_values(["rank", "_co"], na_position="last")
        .drop(columns=["_co"])
        .reset_index(drop=True)
    )
    summary = build_summary(res, stats, sel_date)
    summary = summary[:1] + fmt_regime(gate) + summary[1:]
    # 未接受板块 (被退回) → SUMMARY 顶部醒目标注原因, 清单仍照常输出
    if rejected:
        rej = ["⚠ 以下板块当日未接受 (被退回, 未出股):"]
        rej += [f"    {BOARD_LABEL.get(b, b)}: {r}" for b, r in rejected.items()]
        summary = summary[:1] + rej + summary[1:]
    print(
        f"[enrich] {len(res)} 行 ({(pd.Timestamp.now() - t0).total_seconds():.0f}s)",
        flush=True,
    )

    suffix = _module_suffix(module)
    stamp = str(sel_date.date()).replace("-", "")
    csv_path = STOCK_LIST_DIR / f"parallel_shortlist_{stamp}{suffix}.csv"
    res.to_csv(csv_path, index=False)
    print(f"[saved] {csv_path}", flush=True)
    # WORM: 同名旧文件若被 Word 锁定则换带标记的新名, 不覆盖不丢失
    docx_path = STOCK_LIST_DIR / f"STOCK LIST {stamp}{suffix}.docx"
    try:
        write_docx(res, summary, sel_date, docx_path, merged, module, rejected)
    except PermissionError:
        docx_path = STOCK_LIST_DIR / f"STOCK LIST {stamp}{suffix}_perhorizon.docx"
        write_docx(res, summary, sel_date, docx_path, merged, module, rejected)
        print(f"[warn] 原 docx 被占用, 已写 {docx_path.name}", flush=True)
    if docx_path.exists():
        print(f"[saved] {docx_path}", flush=True)
    xlsx_path = STOCK_LIST_DIR / f"STOCK LIST {stamp}{suffix}.xlsx"
    try:
        write_xlsx(res, summary, sel_date, xlsx_path, merged, module)
    except PermissionError:
        xlsx_path = STOCK_LIST_DIR / f"STOCK LIST {stamp}{suffix}_perhorizon.xlsx"
        try:
            write_xlsx(res, summary, sel_date, xlsx_path, merged, module)
            print(f"[warn] 原 xlsx 被占用, 已写 {xlsx_path.name}", flush=True)
        except PermissionError:
            ts = pd.Timestamp.now().strftime("%H%M%S")
            xlsx_path = (
                STOCK_LIST_DIR / f"STOCK LIST {stamp}{suffix}_perstock_{ts}.xlsx"
            )
            write_xlsx(res, summary, sel_date, xlsx_path, merged, module)
            print(
                f"[warn] 原 xlsx + perhorizon 均被 Excel 占用, 已写 {xlsx_path.name}",
                flush=True,
            )
    if xlsx_path.exists():
        print(f"[saved] {xlsx_path}", flush=True)

    lines = summary[:1] + fmt_merged(merged) + [""] + summary[1:]
    print("\n" + "\n".join(lines), flush=True)
    for board in ("main", "dual"):
        b = res[res["board"] == board]
        if b.empty:
            continue
        print(f"\n══ {board} ══")
        for cut, g in b.groupby("cut", sort=True):
            g = g.sort_values(CAND_RANK_KEY, ascending=False)
            print(f"── {cut} ──")
            for i, (_, r) in enumerate(g.iterrows()):
                tag = "★共现" if r["co_occur"] else "    "
                ex = "  ".join(
                    f"{h}:{r[f'pred_mag_{h}']:+.1%}/{r[f'pred_prob_{h}']:.0%}"
                    for h in HORIZONS
                )
                print(
                    f"  #{i + 1} {tag} {r['symbol']:<8} "
                    f"score={r['score']:.4f} [{ex}] ({r['systems']})"
                )

    if watch:
        print("\n══ 指定个股 ══", flush=True)
        sl = set(res["symbol"])
        for s in watch:
            if s in sl:
                row = res[(res["symbol"] == s) & (res["cut"] == "T-5")].iloc[0]
                seg = "  ".join(
                    f"T+{h[:-1]} {row[f'pred_mag_{h}']:+.1%}({row[f'pred_prob_{h}']:.0%})"
                    for h in HORIZONS
                )
                print(f"{s}: 入选短名单 | 预期涨幅(MFE)/概率: {seg}", flush=True)
            else:
                print(
                    f"{s}: 未入选 — 模型今日未对其打分, 无前瞻预测 "
                    f"(模型只对 top-N 产生预测)",
                    flush=True,
                )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
