"""把 legacy 清单 (list_{trade_date}.parquet) 交付到 STOCK LIST 目录 (WORM).

文件名含 交易日(date) + 模块版本(module) — 回归测试按 module 分组评估各版本表现:
  legacy_stocklist_{date}__{module}.csv   全列 CSV (机器/回归 源真相)
  legacy_stocklist_{date}__{module}.md    Markdown 表格 (可读 + git 可 diff)
  LEGACY STOCK LIST {date}__{module}.docx  Word 表格 (人读)

module 来源: models/pipeline1/current_meta.json (双板同 tag → 单一 tag).
旧文件不覆盖 (WORM).

用法: python scripts/_deliver_legacy_list.py [YYYYMMDD]
"""

import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))

import pandas as pd

from config.settings import STOCK_LIST_DIR
from scripts._stall_marker import stall_marker

try:
    from docx import Document

    HAVE_DOCX = True
except Exception:  # pragma: no cover
    HAVE_DOCX = False

LIST_DIR = "data/lists"


def resolve_module(df: pd.DataFrame, trade_date: str) -> str:
    """模块标识: 优先 current_meta.json; 无则从清单 model_version 列推导."""
    from app.pipeline1.model_meta import load_modules, module_id

    mods = load_modules()
    mid = module_id(mods)
    if mid != "na":
        return mid
    if "model_version" in df.columns:
        uniq = sorted(set(df["model_version"].astype(str)) - {"nan"})
        if len(uniq) == 1:
            return uniq[0]
        if uniq:
            return "__".join(uniq)
    return f"{trade_date}_na"


def _board_reject_reasons(
    cand: pd.DataFrame | None, final: pd.DataFrame
) -> dict[str, str]:
    """各板块被整体退回 (有候选但最终清单 0 只) 的原因 — 重放 E7 可计算闸.

    cand: candidates_{date}.parquet (raw 预测, 含 board/pred_q50_3d/5d/pred_ret_10d).
    final: 最终清单 (list_{date}.parquet). 只在最终清单缺席的板块上推导.
    """
    if cand is None or "board" not in cand.columns:
        return {}
    final_boards = set(final["board"]) if "board" in final.columns else set()
    reasons: dict[str, str] = {}
    for board, sub in cand.groupby("board"):
        if board in final_boards:
            continue
        if {"pred_q50_3d", "pred_q50_5d"} <= set(sub.columns):
            g3 = (sub["pred_q50_3d"] > 0) & (sub["pred_q50_5d"] > 0)
            if not bool(g3.any()):
                reasons[board] = "E7 闸3: 3d/5d 中位数预期均非正 (全灭)"
                continue
        if "pred_ret_10d" in sub.columns:
            g2 = sub["pred_ret_10d"] > 0
            if not bool(g2.any()):
                reasons[board] = "E7 闸2: 10d 净预期非正 (全灭)"
                continue
        reasons[board] = "候选存在但未进最终清单 (排名/行业/风控过滤)"
    return reasons


def write_md(
    df: pd.DataFrame,
    path: str,
    module: str,
    rejected: dict[str, str] | None = None,
) -> None:
    cols = [
        "symbol",
        "board",
        "score",
        "weight",
        "compound_ret",
        "prob_up",
        "prob_up_3d",
        "pred_ret_3d",
        "pred_ret_5d",
        "pred_q50_3d",
        "pred_q50_5d",
        "pain_prob",
        "stall_flag",
        "limit_flag",
        "model_version",
    ]
    cols = [c for c in cols if c in df.columns]
    sub = df[cols].copy()
    for c in sub.columns:
        if sub[c].dtype == "float64":
            sub[c] = sub[c].map(lambda v: f"{v:.4f}")
    sub.insert(0, "rk", range(1, len(sub) + 1))
    base = os.path.basename(path)
    trade_date = base.split("__")[0].replace("legacy_stocklist_", "")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(f"# LEGACY 股票清单 {trade_date} (module {module})\n\n")
        fh.write(
            f"交易日 {trade_date} · module {module} · E7 闸3 = 3d/5d 中位数均正 · "
            f"排序 = d3 目标 (50% d3涨幅 + 50% d3概率 归一化混合) · {len(sub)} 只\n\n"
        )
        # 参与度提示 (2026-08-19): 高基线日模型整体负期望 → 建议降参与
        if "advice" in df.columns and df["advice"].iloc[0]:
            fh.write(df["advice"].iloc[0] + "\n\n")
        # 被整体退回的板块: 仍出清单, 醒目标注未接受原因 (不静默跳过)
        for b, r in (rejected or {}).items():
            fh.write(f"⚠ {b} 未接受 (被退回): {r} — 当日未出股\n\n")
        fh.write(sub.to_markdown(index=False))
        fh.write("\n")


def main():
    trade_date = sys.argv[1] if len(sys.argv) > 1 else "20260805"
    src = os.path.join(LIST_DIR, f"list_{trade_date}.parquet")
    if not os.path.exists(src):
        raise SystemExit(f"无清单: {src}")
    df = pd.read_parquet(src)
    if "symbol" in df.columns:
        df["symbol"] = df["symbol"].astype(str)
    # 滞涨标记 (2026-08-19 用户方案): 入选 + 近10日滞涨<2% + 近20日入选≥3 → 洗盘待爆发
    df = stall_marker(df, trade_date, "legacy_stocklist_")
    module = resolve_module(df, trade_date)
    os.makedirs(str(STOCK_LIST_DIR), exist_ok=True)

    # 被整体退回的板块 (有候选但最终清单 0 只): 仍出该板清单, 醒目标注未接受原因
    cand = None
    cand_path = os.path.join(LIST_DIR, f"candidates_{trade_date}.parquet")
    if os.path.exists(cand_path):
        cand = pd.read_parquet(cand_path)
    rejected = _board_reject_reasons(cand, df)

    csv_path = os.path.join(
        str(STOCK_LIST_DIR), f"legacy_stocklist_{trade_date}__{module}.csv"
    )
    df.to_csv(csv_path, index=False)
    print(f"[csv] {csv_path} ({len(df)} 只)")

    md_path = os.path.join(
        str(STOCK_LIST_DIR), f"legacy_stocklist_{trade_date}__{module}.md"
    )
    write_md(df, md_path, module, rejected)
    print(f"[md] {md_path}")

    for b, r in rejected.items():
        npath = os.path.join(
            str(STOCK_LIST_DIR),
            f"legacy_stocklist_{b}_{trade_date}__{module}_REJECTED.txt",
        )
        with open(npath, "w", encoding="utf-8") as fh:
            fh.write(
                f"LEGACY {b} 清单 {trade_date} (module {module}): "
                f"未接受 (被退回) — {r}\n"
            )
        print(f"[rejected] {b}: {r} → {os.path.basename(npath)}")

    docx_path = ""
    if HAVE_DOCX and len(df):
        doc = Document()
        doc.add_heading(f"LEGACY 股票清单 {trade_date} (module {module})", level=0)
        doc.add_paragraph(
            f"交易日 {trade_date} · module {module} · "
            f"E7 闸3 = 3d/5d 中位数均正 · 排序 = d3 目标 (50% d3涨幅 + 50% d3概率) · {len(df)} 只"
        )
        n_stall = int((df["stall_flag"] != "").sum()) if "stall_flag" in df.columns else 0
        if n_stall:
            doc.add_paragraph(
                f"⚠ 洗盘待爆发 {n_stall} 只 (入选+近10日滞涨<2%+近20日入选≥3, 见 stall_flag 列)",
            )
        for b, r in rejected.items():
            p = doc.add_paragraph()
            run = p.add_run(f"⚠ {b} 未接受 (被退回): {r} — 当日未出股")
            run.bold = True
        cols = [
            "symbol",
            "board",
            "score",
            "weight",
            "compound_ret",
            "prob_up",
            "prob_up_3d",
            "pred_ret_3d",
            "pred_ret_5d",
            "pred_q50_3d",
            "pred_q50_5d",
            "pain_prob",
            "stall_flag",
            "model_version",
        ]
        cols = [c for c in cols if c in df.columns]
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(cols):
            tbl.rows[0].cells[i].text = h
        for _, r in df.iterrows():
            c = tbl.add_row().cells
            for i, col in enumerate(cols):
                v = r[col]
                c[i].text = f"{v:.4f}" if isinstance(v, float) else str(v)
        docx_path = os.path.join(
            str(STOCK_LIST_DIR), f"LEGACY STOCK LIST {trade_date}__{module}.docx"
        )
        doc.save(docx_path)
        print(f"[docx] {docx_path}")

    if not len(df):
        notice = os.path.join(
            str(STOCK_LIST_DIR), f"legacy_list_{trade_date}__{module}.txt"
        )
        with open(notice, "w", encoding="utf-8") as fh:
            fh.write(
                f"LEGACY 清单 {trade_date} (module {module}): 空 (E7 无合格标的)\n"
            )
        print(f"[warn] 空清单 → {notice}")


if __name__ == "__main__":
    main()
