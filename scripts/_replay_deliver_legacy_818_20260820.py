"""_replay_deliver_legacy_818_20260820.py — 8/18 cutoff × 20260820 模块 legacy 清单交付 (2026-08-20).

用户请求: 注意到 legacy raw 文件 (legacy_preds_raw_*__20260820.csv) 存在,
"produce legacy stock list as well" → 用 20260820 模块 (08-20 19:39 重训, 训练含
8/19 涨停标签 = 诊断口径非 PIT) 重放 8/18 cutoff 的完整 legacy 清单并交付.

口径: 与 _diag_300911_818_20260820.py 同款面板切片 (≤8/18, 末 300 交易日),
镜像生产链 daily_pipeline.run() 的 legacy 段:
  clean → 分板特征+预测 (main_20260820/dual_20260820 bundles) → pool_blend_cut
  → compute_scores → entry_filter (E7 生产闸, market_state=range)
  → _rank_by_magnitude (pred_ret_10d 降序, 2026-08-07 定案) → apply_industry_limit
  → head(15) (TOP_N, 与 8/18 实发 15 只同档).
滞涨标记: stall_marker 传 ≤8/18 切片面板 (临时 parquet) → base_rate 为 8/18 当日口径
  (stall_marker 默认读全量面板会偏 8/19-8/20 未来数据).

交付 (WORM, 不覆盖既有 20260818/neg200 文件):
  legacy_stocklist_20260818__20260820.{csv,md} + LEGACY STOCK LIST 20260818__20260820.docx
输出: BACKTEST_RESULT_DIR/_replay_818_20260820_<ts>/report.json.
用法: python scripts/_replay_deliver_legacy_818_20260820.py
"""

from __future__ import annotations

import gc
import json
import os
import sys
import time
from pathlib import Path

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))

import pandas as pd
import pyarrow.dataset as ds

from app.pipeline1.cleaning_pipeline import CleaningPipeline
from app.pipeline1.feature_engine_v35 import FeatureEngineV35
from app.pipeline1.list_generator import ListGenerator
from app.pipeline1.predictor import V35Predictor
from config.settings import BACKTEST_RESULT_DIR, PANEL_V3_PATH, STOCK_LIST_DIR
from scripts._stall_marker import stall_marker

TARGET = "300911"
TRADE_DATE = pd.Timestamp("2026-08-18")
TRADE_STR = "20260818"
MODULE = "20260820"
BUNDLES = {
    "main": "models/pipeline1/main_20260820.pkl",
    "dual": "models/pipeline1/dual_20260820.pkl",
}
TAIL_DAYS = 300  # 与 _gen_legacy_list.py 同款内存切片 (特征等价已验证)
TOP_N = 15  # 生产 TOP_N (D18 正常日 cap>=1.0)
MARGIN_DUAL = 0.08  # LEGACY_ENTRY_GATE.prob_margin dual
PAIN_MAX_DUAL = 0.4  # LEGACY_ENTRY_GATE.pain_max dual

report: dict = {"cutoff": str(TRADE_DATE.date()), "module": MODULE, "bundles": BUNDLES}


def log(msg: str) -> None:
    print(msg, flush=True)


def ram_gb() -> float:
    try:
        import psutil

        return psutil.virtual_memory().available / 1e9
    except Exception:
        return -1.0


def board_reject_reasons(cand: pd.DataFrame, final: pd.DataFrame) -> dict[str, str]:
    """镜像 _deliver_legacy_list._board_reject_reasons (被整体退回的板块)."""
    final_boards = set(final["board"]) if "board" in final.columns else set()
    reasons: dict[str, str] = {}
    for board, sub in cand.groupby("board"):
        if board in final_boards:
            continue
        if {"pred_q50_3d", "pred_q50_5d"} <= set(sub.columns):
            g3 = (sub["pred_q50_3d"] > 0) & (sub["pred_q50_5d"] > 0)
            if not bool(g3.any()):
                reasons[board] = "E7 闸3: 3d/5d 中位数预期均非正 (全灭)"
                continue
        if "pred_ret_10d" in sub.columns:
            g2 = sub["pred_ret_10d"] > 0
            if not bool(g2.any()):
                reasons[board] = "E7 闸2: 10d 净预期非正 (全灭)"
                continue
        reasons[board] = "候选存在但未进最终清单 (排名/行业/风控过滤)"
    return reasons


def write_md(df: pd.DataFrame, path: str, rejected: dict[str, str]) -> None:
    cols = [
        "symbol",
        "board",
        "score",
        "weight",
        "compound_ret",
        "prob_up",
        "prob_up_3d",
        "pred_ret_3d",
        "pred_ret_5d",
        "pred_q50_3d",
        "pred_q50_5d",
        "pain_prob",
        "stall_flag",
        "limit_flag",
        "model_version",
    ]
    cols = [c for c in cols if c in df.columns]
    sub = df[cols].copy()
    for c in sub.columns:
        if sub[c].dtype == "float64":
            sub[c] = sub[c].map(lambda v: f"{v:.4f}")
    sub.insert(0, "rk", range(1, len(sub) + 1))
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(
            f"# LEGACY 股票清单 {TRADE_STR} (module {MODULE}) — 重放 (诊断口径)\n\n"
        )
        fh.write(
            f"交易日 {TRADE_STR} · module {MODULE} (08-20 重训, 训练含 8/19 标签非 PIT) · "
            f"E7 闸3 = 3d/5d 中位数均正 · 排序 = pred_ret_10d 降序 "
            f"(生产 _rank_by_magnitude 口径) · {len(df)} 只\n\n"
        )
        if "advice" in df.columns and len(df) and df["advice"].iloc[0]:
            fh.write(df["advice"].iloc[0] + "\n\n")
        for b, r in rejected.items():
            fh.write(f"⚠ {b} 未接受 (被退回): {r} — 当日未出股\n\n")
        fh.write(sub.to_markdown(index=False))
        fh.write("\n")


def write_docx(df: pd.DataFrame, path: str, rejected: dict[str, str]) -> bool:
    try:
        from docx import Document
    except Exception:
        return False
    doc = Document()
    doc.add_heading(
        f"LEGACY 股票清单 {TRADE_STR} (module {MODULE}) — 重放 (诊断口径)", level=0
    )
    doc.add_paragraph(
        f"交易日 {TRADE_STR} · module {MODULE} (08-20 重训, 训练含 8/19 标签非 PIT) · "
        f"E7 闸3 = 3d/5d 中位数均正 · 排序 = pred_ret_10d 降序 · {len(df)} 只"
    )
    n_stall = int((df["stall_flag"] != "").sum()) if "stall_flag" in df.columns else 0
    if n_stall:
        doc.add_paragraph(
            f"⚠ 洗盘待爆发 {n_stall} 只 (入选+近10日滞涨<2%+近20日入选≥3, 见 stall_flag 列)",
        )
    for b, r in rejected.items():
        p = doc.add_paragraph()
        run = p.add_run(f"⚠ {b} 未接受 (被退回): {r} — 当日未出股")
        run.bold = True
    cols = [
        "symbol",
        "board",
        "score",
        "weight",
        "compound_ret",
        "prob_up",
        "prob_up_3d",
        "pred_ret_3d",
        "pred_ret_5d",
        "pred_q50_3d",
        "pred_q50_5d",
        "pain_prob",
        "stall_flag",
        "model_version",
    ]
    cols = [c for c in cols if c in df.columns]
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(cols):
        tbl.rows[0].cells[i].text = h
    for _, r in df.iterrows():
        c = tbl.add_row().cells
        for i, col in enumerate(cols):
            v = r[col]
            c[i].text = f"{v:.4f}" if isinstance(v, float) else str(v)
    doc.save(path)
    return True


def main() -> int:
    t0 = time.time()
    log(f"[ram] available {ram_gb():.1f} GB")
    if ram_gb() < 2.5:
        log("[FATAL] 内存 < 2.5GB, 拒绝启动 (重训/其他重活占用?)")
        return 3

    # ---------- 1. 面板: 全市场 ≤ 8/18, 末 300 交易日 ----------
    log("[1] 读面板 (≤8/18, 末300交易日) ...")
    d = ds.dataset(str(PANEL_V3_PATH))
    df = d.to_table(filter=ds.field("date") <= TRADE_DATE).to_pandas()
    df["symbol"] = df["symbol"].astype(str)
    df["date"] = pd.to_datetime(df["date"]).dt.normalize()
    dates = sorted(df["date"].unique())
    cut = dates[-TAIL_DAYS]
    df = df[df["date"] >= cut].sort_values(["symbol", "date"]).reset_index(drop=True)
    log(f"[1] {len(df):,} 行 {cut.date()}..{dates[-1].date()}")
    report["panel_rows"] = int(len(df))
    report["panel_start"] = str(cut.date())

    # ---------- 2. 清洗 (生产口径, pool_blend=True) ----------
    cleaner = CleaningPipeline()
    main_df, dual_df, _state = cleaner.run_inference(df, pool_blend=True)
    del df
    gc.collect()
    for b, sub in (("main", main_df), ("dual", dual_df)):
        day = sub[pd.to_datetime(sub["date"]) == TRADE_DATE]
        log(f"[2] {b} 清洗后 {len(sub):,} 行; 8/18 截面 {len(day):,} 只")
        report[f"pool_{b}_818"] = int(len(day))

    # ---------- 3. 分板特征 + 预测 (20260820 模块) ----------
    predictor = V35Predictor(BUNDLES)
    frames = []
    for board, d in (("main", main_df), ("dual", dual_df)):
        feat = FeatureEngineV35().build(
            d,
            None,
            inference_cols=predictor.bundles[board]["feature_cols"],
            cross_sectional_rank=True,
        )
        log(f"[3] {board} 特征帧 {len(feat):,} 行 × {len(feat.columns)} 列")
        today_feat = feat[feat["date"] == feat["date"].max()].copy()
        pred = predictor.predict(today_feat, board)
        log(f"[3] {board} 预测 {len(pred)} 只")
        frames.append(pred)
        del feat, today_feat, pred, d
        gc.collect()
    candidates = pd.concat(frames, ignore_index=True)
    report["n_candidates"] = int(len(candidates))
    report["candidates_by_board"] = candidates["board"].value_counts().to_dict()

    # ---------- 4. pool_blend_cut (生产链原样调用; 08-20 实证 fail-open) ----------
    has_ls = "liquidity_score" in candidates.columns
    log(
        f"[4] candidates 含 liquidity_score: {has_ls} → blend 切池 {'生效' if has_ls else 'FAIL-OPEN 不切'}"
    )
    report["blend_cut_active"] = bool(has_ls)
    candidates = cleaner.pool_blend_cut(candidates)
    report["n_after_blend_cut"] = int(len(candidates))

    # ---------- 5. E7 生产闸 + 排名 + 行业限制 ----------
    lister = ListGenerator()
    scored = lister.compute_scores(candidates)
    passed = lister.entry_filter(scored, market_state="range")
    ranked = lister._rank_by_magnitude(passed)
    final = lister.apply_industry_limit(ranked).reset_index(drop=True).head(TOP_N)
    log(
        f"[5] E7 过闸 {len(passed)} 只 (main={int((passed['board'] == 'main').sum())} / "
        f"dual={int((passed['board'] != 'main').sum())}); 行业限制+TOP15 后 {len(final)} 只"
    )
    report["n_passed"] = int(len(passed))
    report["passed_by_board"] = passed["board"].value_counts().to_dict()
    report["n_final"] = int(len(final))

    # ---------- 6. 300911 全链路判定 ----------
    r = scored[scored["symbol"] == TARGET]
    if len(r):
        x = r.iloc[0]
        base = x["base_rate"]
        checks = {
            "闸1 compound_prob > base_rate+0.08": (
                float(x["compound_prob"]),
                float(base + MARGIN_DUAL),
                bool(x["compound_prob"] > base + MARGIN_DUAL),
            ),
            "闸2 pred_ret_10d > 0": (
                float(x["pred_ret_10d"]),
                0.0,
                bool(x["pred_ret_10d"] > 0),
            ),
            "闸3a pred_q50_3d > 0": (
                float(x["pred_q50_3d"]),
                0.0,
                bool(x["pred_q50_3d"] > 0),
            ),
            "闸3b pred_q50_5d > 0": (
                float(x["pred_q50_5d"]),
                0.0,
                bool(x["pred_q50_5d"] > 0),
            ),
            "E2  pain_prob <= 0.4": (
                float(x["pain_prob"]),
                PAIN_MAX_DUAL,
                bool(x["pain_prob"] <= PAIN_MAX_DUAL),
            ),
        }
        log("\n== 300911 8/18 (20260820 模块) 闸门逐项 ==")
        for k, (v, th, ok) in checks.items():
            log(f"  {k}: {v:+.4f} vs {th:+.4f} → {'PASS' if ok else 'FAIL'}")
        report["gates"] = {
            k: {"value": v, "threshold": th, "pass": ok}
            for k, (v, th, ok) in checks.items()
        }
        in_list = TARGET in set(final["symbol"])
        log(
            f"\n  → 300911 最终{'入选' if in_list else '未入选'} 8/18 清单 (20260820 模块)"
        )
        report["in_list"] = bool(in_list)
    else:
        log(f"[6] {TARGET} 不在预测输出 (被清洗层剔除)")
        report["in_list"] = False
        report["gates"] = None

    # ---------- 7. 滞涨标记 (8/18 当日 base_rate 口径) ----------
    tmp_pq = Path(BACKTEST_RESULT_DIR) / "_replay_818_20260820_tmp_panel.pq"
    tmp_pq.parent.mkdir(parents=True, exist_ok=True)
    # 重读 ≤8/18 面板 5 列 (stall_marker 需要完整窗口算 base_rate)
    d5 = ds.dataset(str(PANEL_V3_PATH))
    p5 = d5.to_table(
        columns=["symbol", "date", "close_hfq", "high_hfq", "amount", "board"],
        filter=ds.field("date") <= TRADE_DATE,
    ).to_pandas()
    p5.to_parquet(tmp_pq, index=False)
    del p5
    gc.collect()
    marked = stall_marker(final, TRADE_STR, "legacy_stocklist_", panel_path=tmp_pq)
    tmp_pq.unlink(missing_ok=True)
    n_stall = int((marked["stall_flag"] != "").sum())
    n_lim = int((marked["limit_flag"] != "").sum())
    log(
        f"[7] stall_marker: stall={n_stall} 只, limit={n_lim} 只, "
        f"market_base_rate={marked['market_base_rate'].iloc[0] if len(marked) else None}"
    )
    report["n_stall"] = n_stall
    report["n_limit"] = n_lim
    report["market_base_rate"] = (
        float(marked["market_base_rate"].iloc[0]) if len(marked) else None
    )

    # ---------- 8. 交付 (WORM) ----------
    mod = "20260820"  # bundle 显式指定, 不依赖 current_meta
    if "model_version" not in marked.columns:
        marked["model_version"] = mod
    rejected = board_reject_reasons(candidates, marked)
    csv_path = str(STOCK_LIST_DIR / f"legacy_stocklist_{TRADE_STR}__{mod}.csv")
    marked.to_csv(csv_path, index=False)
    log(f"[8] [csv] {csv_path} ({len(marked)} 只)")
    md_path = str(STOCK_LIST_DIR / f"legacy_stocklist_{TRADE_STR}__{mod}.md")
    write_md(marked, md_path, rejected)
    log(f"[8] [md] {md_path}")
    docx_path = str(STOCK_LIST_DIR / f"LEGACY STOCK LIST {TRADE_STR}__{mod}.docx")
    if write_docx(marked, docx_path, rejected):
        log(f"[8] [docx] {docx_path}")
    for b, r in rejected.items():
        npath = str(
            STOCK_LIST_DIR / f"legacy_stocklist_{b}_{TRADE_STR}__{mod}_REJECTED.txt"
        )
        with open(npath, "w", encoding="utf-8") as fh:
            fh.write(
                f"LEGACY {b} 清单 {TRADE_STR} (module {mod}): 未接受 (被退回) — {r}\n"
            )
        log(f"[8] [rejected] {b}: {r}")

    # ---------- 落盘 ----------
    out_dir = Path(BACKTEST_RESULT_DIR) / (
        "_replay_818_20260820_" + time.strftime("%Y%m%d_%H%M%S")
    )
    out_dir.mkdir(parents=True, exist_ok=True)
    with open(out_dir / "report.json", "w", encoding="utf-8") as fh:
        json.dump(report, fh, ensure_ascii=False, indent=2, default=str)
    log(f"\n[done] WORM -> {out_dir / 'report.json'} ({time.time() - t0:.0f}s)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
