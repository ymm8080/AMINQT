"""Tests for scripts/_shortlist_t5_t10 交付落盘的 stall_flag 容错 (2026-08-20).

背景: 20:12 自动化 deliver_parallel 崩溃 — write_xlsx 里 m[mcols] 抛 KeyError
"['stall_flag'] not in index": merged 帧无 stall_flag 列 (stall_marker 只作用于 res,
main() 在 build_merged 之后才调用).
修复语义 = 给 merged 补空列 ("" 而非 NaN) + docx 侧 r.get("stall_flag", "") 同容错;
res 经 stall_marker 后必有真实 stall_flag, 原样写出.
"""

import pandas as pd

from scripts._shortlist_t5_t10 import write_docx, write_xlsx

HORIZONS = ("3d", "5d", "10d")
PRED_COLS = [f"{k}_{h}" for h in HORIZONS for k in ("pred_mag", "pred_prob")]

SEL_DATE = pd.Timestamp("2026-02-05")


def _res_df(stall: str = "") -> pd.DataFrame:
    """最小 res 帧: 覆盖 write_docx 分板 sheet / write_xlsx 数据 sheet 的全部列."""
    return pd.DataFrame(
        [
            {
                "date": "2026-02-05",
                "board": "main",
                "cut": "T-10",
                "rank": 1,
                "symbol": "600001",
                "systems": "M2026",
                "co_occur": True,
                "score": 0.5,
                "score_w": 0.25,
                "过门": "是",
                "stall_flag": stall,
                **{c: 0.05 if c.startswith("pred_mag") else 0.6 for c in PRED_COLS},
            }
        ]
    )


def _merged_df(stall=None) -> pd.DataFrame:
    """最小 merged 帧: 覆盖 docx/xlsx 合并排名 sheet 的全部列.

    stall=None → 无 stall_flag 列 (崩溃现场形态); 否则为给定值 (含真实值场景).
    """
    d = {
        "rank": 1,
        "symbol": "600001",
        "board": "main",
        "systems": "M2026",
        "score": 0.5,
        "score_w": 0.25,
        "in_t5": True,
        "过门": "是",
        **{c: 0.05 if c.startswith("pred_mag") else 0.6 for c in PRED_COLS},
    }
    if stall is not None:
        d["stall_flag"] = stall
    return pd.DataFrame([d])


# ---------- write_xlsx ----------


def test_xlsx_merged_without_stall_flag_writes_empty(tmp_path):
    """崩溃现场: merged 无 stall_flag → 不抛异常, 合并排名 sheet 该列为空串."""
    res = _res_df(stall="洗盘待爆发")
    merged = _merged_df(stall=None)
    out = tmp_path / "out.xlsx"
    write_xlsx(res, [], SEL_DATE, out, merged, "test")
    from openpyxl import load_workbook

    wb = load_workbook(out)
    ws = wb["合并排名"]
    hdr = [c.value for c in ws[1]]
    assert "stall_flag" in hdr
    assert (
        hdr
        == [
            "rank",
            "symbol",
            "board",
            "module",
            "score",
            "score_w",
            "in_t5",
            "过门",
            "stall_flag",
        ]
        + PRED_COLS
    )
    j = hdr.index("stall_flag") + 1  # openpyxl 1-based
    assert ws.cell(row=2, column=j).value in ("", None)


def test_xlsx_merged_with_stall_flag_writes_real_value(tmp_path):
    """merged 自带 stall_flag → 不被补空列覆盖, 原值写出."""
    merged = _merged_df(stall="洗盘待爆发")
    out = tmp_path / "out.xlsx"
    write_xlsx(_res_df(), [], SEL_DATE, out, merged, "test")
    from openpyxl import load_workbook

    ws = load_workbook(out)["合并排名"]
    hdr = [c.value for c in ws[1]]
    j = hdr.index("stall_flag") + 1
    assert ws.cell(row=2, column=j).value == "洗盘待爆发"


def test_xlsx_res_stall_flag_writes_real_value(tmp_path):
    """res 经 stall_marker 后有真实 stall_flag → 数据 sheet 原样写出."""
    out = tmp_path / "out.xlsx"
    write_xlsx(_res_df(stall="洗盘待爆发"), [], SEL_DATE, out, None, "test")
    from openpyxl import load_workbook

    ws = load_workbook(out)["短名单 T-10"]
    hdr = [c.value for c in ws[1]]
    assert (
        hdr
        == [
            "date",
            "board",
            "cut",
            "rank",
            "symbol",
            "module",
            "co_occur",
            "score",
            "过门",
            "stall_flag",
        ]
        + PRED_COLS
    )
    j = hdr.index("stall_flag") + 1
    assert ws.cell(row=2, column=j).value == "洗盘待爆发"


# ---------- write_docx ----------


def test_docx_merged_without_stall_flag_tolerated(tmp_path):
    """merged/res 均无 stall_flag → r.get 容错, 表头含 stall_flag 且数据行为空串."""
    out = tmp_path / "out.docx"
    write_docx(_res_df(stall=""), [], SEL_DATE, out, _merged_df(stall=None), "test", {})
    from docx import Document

    doc = Document(out)
    t_merged = doc.tables[0]
    hdr = [c.text for c in t_merged.rows[0].cells]
    assert (
        hdr
        == [
            "rank",
            "symbol",
            "board",
            "module",
            "score",
            "score_w",
            "in_t5",
            "过门",
            "stall_flag",
        ]
        + PRED_COLS
    )
    assert len(hdr) == len(t_merged.rows[1].cells)  # 单元格索引与 mcols 长度一致
    assert t_merged.rows[1].cells[8].text == ""
    t_board = doc.tables[1]  # 分板 sheet
    hdr_b = [c.text for c in t_board.rows[0].cells]
    assert hdr_b[7] == "stall_flag"
    assert len(hdr_b) == len(t_board.rows[1].cells)
    assert t_board.rows[1].cells[7].text == ""


def test_docx_stall_flag_real_value_written(tmp_path):
    """res/merged 有 stall_flag → 两表写出真实值 (docx 固定下标与列序对齐)."""
    out = tmp_path / "out.docx"
    write_docx(
        _res_df(stall="洗盘待爆发"),
        [],
        SEL_DATE,
        out,
        _merged_df(stall="洗盘待爆发"),
        "test",
        {},
    )
    from docx import Document

    doc = Document(out)
    t_merged = doc.tables[0]
    assert t_merged.rows[1].cells[8].text == "洗盘待爆发"
    assert t_merged.rows[1].cells[9].text == "+5.0%"  # pred_mag_3d 紧随其后
    t_board = doc.tables[1]
    assert t_board.rows[1].cells[7].text == "洗盘待爆发"
    assert t_board.rows[1].cells[8].text == "+5.0%"  # pred_mag_3d 紧随其后
