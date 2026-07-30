# -*- coding: utf-8 -*-
"""Convert FEATURE ADOPTION ANALYSIS markdown to Word document."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

MD = "data/feature_logs/FEATURE ADOPTION ANALYSIS 20260729"
OUT = "reference/FEATURE REVIEW/FEATURE ADOPTION ANALYSIS 20260729.docx"

doc = Document()

# -- Styles --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p

def add_table_from_md(lines_iter):
    """Consume markdown table lines from iterator, return docx table."""
    rows = []
    for line in lines_iter:
        line = line.strip()
        if not line.startswith('|'):
            break
        # skip separator rows like |---|----|
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows.append(cells)
    if not rows:
        return None, None

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < ncols:
                cell = table.cell(i, j)
                cell.text = cell_text
                # bold header
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(9)
                else:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)

    return table, rows

# -- Read and convert --
with open(MD, 'r', encoding='utf-8') as fh:
    lines = fh.readlines()

it = iter(lines)
in_table = False
table_lines = []

for line in it:
    line_rstrip = line.rstrip()

    # Blank line
    if not line_rstrip:
        if in_table:
            add_table_from_md(iter(table_lines))
            in_table = False
            table_lines = []
        continue

    # Headings
    if line_rstrip.startswith('# '):
        add_heading(line_rstrip[2:], level=1)
        continue
    if line_rstrip.startswith('## '):
        add_heading(line_rstrip[3:], level=2)
        continue
    if line_rstrip.startswith('### '):
        add_heading(line_rstrip[4:], level=3)
        continue

    # Horizontal rule
    if line_rstrip.strip() == '---':
        doc.add_paragraph('_' * 60)
        continue

    # Table detection
    if line_rstrip.startswith('|'):
        table_lines.append(line_rstrip)
        in_table = True
        continue

    # Code block
    if line_rstrip.startswith('```'):
        continue

    # Regular paragraph
    # Strip markdown bold markers
    text = line_rstrip
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    add_para(text)

# Flush remaining table
if table_lines:
    add_table_from_md(iter(table_lines))

# -- Set narrow margins --
for section in doc.sections:
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

doc.save(OUT)
print(f"Saved: {OUT}")
